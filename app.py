### pyinstaller --hiddenimport win32timezone --onefile -w app.py ---> For recompiling into exe

from tkinter import * ### module for creating the UI
from tkinter.filedialog import askopenfilename ### module for browse button of the app
import tkinter as tk ### module for creating the UI
import sys  ### module for redirecting terminal logs and prints into the tkinter textbox
import threading ### module for redirecting terminal logs and prints into the tkinter textbox
import traceback ### module for redirecting terminal logs and prints into the tkinter textbox
import os ### module for easy access of directories
from PIL import Image, ImageTk ### module for app logo

output_path = os.getcwd()
os.chdir(output_path)

# Color themes
ThemeYellow = '#ffe600'
ThemeGrey = '#333333'
ThemeGrey2 = '#cccccc'

# Read user inputs
file = open(output_path + '/components/UserInputs.txt', 'r')
read = file.readlines()
readSplit = read[0].split(',')

# Setting up variables
Preparer = readSplit[0]
Reviewer = readSplit[1]
fundRequestVal = readSplit[2]

# Display of fund request path
try:
    fundRequestDisplay = '../'+fundRequestVal.split("/")[-3]+'/'+fundRequestVal.split("/")[-2]+'/'+fundRequestVal.split("/")[-1]
except:
    fundRequestDisplay = readSplit[2]

# Remove blurriness in  the app display
import ctypes
ctypes.windll.shcore.SetProcessDpiAwareness(1)

# Setting root
root = tk.Tk()
root.title('WAMapps Automated Memo Tool')

# canvas size
canvas = tk.Canvas(root, width=670, height=700, bg=eyThemeGrey, bd=-2)
canvas.grid(columnspan=3, rowspan=7)

#logo
logo = Image.open('components/logo.JPG')
logo= logo.resize((370, 120))
logo = ImageTk.PhotoImage(logo)
logo_label = tk.Label(image=logo, bd=-2)
logo_label.image = logo
logo_label.grid(columnspan=3, column=0, row=0)

# Input field names
fundRequest_text = Label(text = "        Fund Request * ",bg=eyThemeGrey, fg='white')
Preparer_text = Label(text = "        Preparer * ",bg=eyThemeGrey, fg='white')
Reviewer_text = Label(text = "        Reviewer * ",bg=eyThemeGrey, fg='white')
fundRequest_text.grid(column=0,row=1,sticky = W)
Preparer_text.grid(column=0,row=2,sticky = W)
Reviewer_text.grid(column=0,row=3,sticky = W)

# Input field entries
firstname = StringVar()
lastname = StringVar()
age = StringVar()
fundRequest_entry = Entry(textvariable = age, width = "47")
Preparer_entry = Entry(textvariable = firstname, width = "47")
Reviewer_entry = Entry(textvariable = lastname, width = "47")
Preparer_entry.insert(0, Preparer)
Reviewer_entry.insert(0, Reviewer)
fundRequest_entry.insert(0, fundRequestDisplay)
fundRequest_entry.bind("<Key>", lambda e: "break")
Preparer_entry.grid(column=1,row=2,sticky = W,ipady=4)
Reviewer_entry.grid(column=1,row=3,sticky = W,ipady=4)
fundRequest_entry.grid(column=1,row=1,sticky = W,ipady=4)

# browse button
def open_file():
    browse_text.set("Loading...")
    file = askopenfilename(parent=root, title="Choose a file", filetypes=[("csv file", "*.csv")])
    
    if file =='':
        pass
    else:
        fundRequest_entry.delete(0, END)
        fundRequestDisplay = '../'+file.split("/")[-3]+'/'+file.split("/")[-2]+'/'+file.split("/")[-1]
        fundRequest_entry.insert(0,fundRequestDisplay)
        open_file.file = file
    browse_text.set("Browse")

browse_text = tk.StringVar()
browse_btn = tk.Button(root, textvariable=browse_text, command=open_file, bg=eyThemeYellow, width=8,font="Arial 8 bold" )
browse_text.set("Browse")
browse_btn.grid(column=2, row=1,sticky=W)

# For redirecting of console to tkinter textbox (e.g. tracebacks, prints, etc.)
class Redirect():

    def __init__(self, widget, autoscroll=True):
        self.widget = widget
        self.autoscroll = autoscroll

    def write(self, text):
        self.widget.insert('end', text)
        if self.autoscroll:
            self.widget.see("end")  # autoscroll

# Threading
def run_threading():
    threading.Thread(target=run).start()

# Run button
def run():
  Preparer_info = firstname.get()
  Reviewer_info = lastname.get()
  try:
    fundRequest_info = open_file.file
  except:
    fundRequest_info = fundRequestVal
  text_file = open(output_path + "/components/UserInputs.txt", "w")
  text_file.write(Preparer_info + ',' + Reviewer_info + ',' + fundRequest_info)
  text_file.close()
  main_script()

#Run button
Run_text = tk.StringVar()
Run_btn = tk.Button(root, textvariable=Run_text, command=run_threading, font="Arial 15 bold", bg=eyThemeYellow, height=1, width=15,)
Run_text.set("Run")
Run_btn.grid(column=0, row=4, columnspan=3)

#text box
text_box = tk.Text(root,height=15,width=85, font="Arial 8", bg=eyThemeGrey2)
text_box.grid(column=0, row=5, columnspan=3)

# Redirecting console logs to tkinter textbox
old_stdout = sys.stdout    
sys.stdout = Redirect(text_box)

issues_counter = 0
def main_script():
  try: 
    import time ###to delay the screenshot
    print('\n- - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -')
    print('\nApp will start in a while . . .')
    time.sleep(4)
    print('\nApp started running.')
    print('\nImporting modules . . .')

    import os ### module for easy access of directories
    import pandas as pd ###processing modules
    import docx ### module to fill up word documents
    import numpy as np  ### processing modules
    from datetime import datetime ### module to set datetime format
    from docx.shared import Pt ###module to set formatting of docx
    from docx.shared import Inches ###module to resize images in word
    from docx.shared import RGBColor ###set fontcolor in word
    import pyautogui ###module to set screenshots 
    import re # to delete special characters
    from win32com import client #module to open excel via Python
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT #module to align the pictures
    from PIL import Image, ImageOps #Edited module for adding picture border
    import shutil ### for deleting a folder

    #Hide future warnings
    import warnings
    warnings.simplefilter(action='ignore', category=FutureWarning)

    os.chdir(os.getcwd())

    # Read user inputs
    file = open(output_path + '/components/UserInputs.txt', 'r')
    read = file.readlines()
    readSplit = read[0].split(',')

    # Setting up variables
    Preparer = readSplit[0]
    Reviewer = readSplit[1]
    fundRequest = readSplit[2]
    row_numbers = ['All']  

    #Validation of user inputs
    invalidCount = 0

    #validation - if fund request is not found in the provided link
    try:
        doc_fr = pd.read_csv(fundRequest) 
    except:
        print('\nInvalid Run: Fund Request not found. Could not read file: "' + fundRequest +'". Please check then browse again if necessary')
        invalidCount += 1
        
    ###Code proper
    if invalidCount > 0:
        print('\nRunning Completed: Please fix all the issues noted above then rerun to proceed.')
    else:
        runNum = 0

        os.chdir(output_path)
        #Assign the row number/s to work with
        rowNumbersFinal = []
        for i in range(0,len(doc_fr)):
            rowNumbersFinal.append(i)

        Preparer = Preparer + '\n(WAMApps Preparer)'
        Reviewer = Reviewer + '\n(WAMApps Reviewer)'

        for row_num in rowNumbersFinal:
            print("\nProcessing FR Row Number " + str(row_num)+ ". . .")
            start_time = time.time()
            global issues_counter
            issues_counter = 0
            os.chdir(output_path)
            runNum += 1
            #Functions Section and Global Variables

            header_no = 10

            # #For deleting user defined variables to avoid bugs in bulk processing
            # Global_var = []
            # for var in dir():
            #     Global_var.append(var)

            #For error handling
            class trialContextManager:
                def __enter__(self): pass
                def __exit__(self, *args): return True
            trial = trialContextManager()

            #Screenshot Border Function
            def add_border(input_image, output_image):
                bimg = ImageOps.expand(input_image, border=1,fill='black')
                bimg.save(output_image)

            #Memo Issues and Screenshot
            def memo_final(memo_text,screenshot,reco_text):
                global issues_counter
                issues_counter = issues_counter + 1
                sec3 = docMemo.tables[1].cell(3,2).add_paragraph(style='List Bullet')
                sec3.add_run().add_text(memo_text)
                sec3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                docMemo.tables[1].cell(3,2).add_paragraph().add_run().add_break()
                p = docMemo.tables[2].cell(1,0).add_paragraph(style='List Bullet')
                r = p.add_run()
                r.add_text(memo_text)
                x = docMemo.tables[2].cell(1,0).add_paragraph()
                x.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                x.add_run().add_break()
                for (a,b) in screenshot:
                    x.add_run().add_picture(a,b)
                x.add_run().add_break()
                docMemo.tables[2].cell(1,0).add_paragraph(style='List Paragraph').add_run('Recommendation:')
                reco = docMemo.tables[2].cell(1,0).add_paragraph(style='List Paragraph').add_run(reco_text)
                reco.bold = False
                reco.add_break()

            #Memo Recommendation
            dummy = []
            def recommendation(inputs):
                reco_else = []
                for (value,names,workpaper,newcol,newcol_reclass,sheets,tab,reco,reclass_tab,reco_pic) in inputs:
                    if abs(value)>1:
                        try:
                            names.append(workpaper)
                        except:
                            pass
                        #for recommendation
                        Leads_TB[newcol] = Leads_TB['Current_Year'].apply(lambda x: np.nan if x == 0 else abs(x)- abs(value))
                        Leads_TBClass[newcol] = Leads_TBClass['Current_Year'].apply(lambda x: np.nan if x == 0 else abs(x)- abs(value))
                        Leads_TB[newcol_reclass] = Leads_TB[newcol].apply(lambda x: 'Reclass' if abs(x)<1 else 'No Reclass')
                        Leads_TBClass[newcol_reclass] = Leads_TBClass[newcol].apply(lambda x: 'Reclass' if abs(x)<1 else 'No Reclass')
                        if tab not in sheets.sheet_names:
                            reco.append('The difference in %s is either due to the missing %s input file or incorrect account mapping. As such audit team shall obtain the missing file or modify the account mapping and ask WAMapps team for a refreshed output.'%(workpaper,workpaper))  
                        elif (Leads_TB[newcol_reclass]=='Reclass').any():
                            account_name = Leads_TB[Leads_TB[newcol_reclass]=='Reclass']['Account_Name'].iloc[0]
                            account_class = Leads_TB[Leads_TB[newcol_reclass]=='Reclass']['Class'].iloc[0]
                            reclass_tab = reclass_tab.append(Leads_TB[Leads_TB[newcol_reclass]=='Reclass']['Tab'].iloc[0])
                            account_names_append.append(account_name)
                            account_class_append.append(account_class)
                            reco.append('The difference in %s is due to account: %s. As such the audit team may opt to manually reclassify the said account from %s to its appropriate account class.'%(workpaper,account_name,account_class))
                            if re.sub('[^A-Za-z0-9]+', '',account_class) in class2:
                                reco_pic.append(re.sub('[^A-Za-z0-9]+', '',account_class)+'.png')
                        elif (Leads_TBClass[newcol_reclass]=='Reclass').any():
                            account_class = Leads_TBClass[Leads_TBClass[newcol_reclass]=='Reclass'].index.values.tolist()[0][0]
                            reclass_tab = reclass_tab.append(Leads_TBClass[Leads_TBClass[newcol_reclass]=='Reclass'].index.values.tolist()[0][1])
                            account_class_append.append(account_class)
                            reco.append('The difference in %s is due to account class: %s. As such the audit team may opt to manually reclassify the said account class.'%(workpaper,account_class))
                            if re.sub('[^A-Za-z0-9]+', '',account_class) in class2:
                                reco_pic.append(re.sub('[^A-Za-z0-9]+', '',account_class)+'.png')
                        else:
                            reco_else.append(workpaper)
                if len(reco_else) > 0:
                    reco_else_str = ', '.join(reco_else)
                    reco.append('For the %s variance, the audit team shall investigate the input files of the said accounts and determine the root cause of the variance. They may also opt to reclassify the mapping of certain accounts as this may also be another reason for the variance.'%(reco_else_str))
                else:
                    pass

            #Memo Refer to Tab for Variance
            def refer_memo(inputs):
                for (value,tab,word) in inputs:
                    if len(value)==1:
                        tab_str = ''.join(tab)
                        pic = docMemo.tables[2].cell(1,0).add_paragraph(style='List Paragraph').add_run('**Refer to %s tab of Leads for the %s variance' %(tab_str,word))
                        pic.font.size = Pt(8)
                        pic.italic = True
                        pic.add_picture(value[0],width=Inches(4.5))


            #Memo Completeness Check
            def completeness(memo_text,reco_text):
                sec2a = docMemo.tables[1].cell(2,2).paragraphs[0]
                sec2a.text = (memo_text)
                sec2a.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                sec2a.style = docMemo.styles['List Bullet']
                docMemo.tables[2].cell(1,0).paragraphs[1].add_run().add_break()
                sec3a = docMemo.tables[2].cell(1,0).paragraphs[1]
                sec3a.style = docMemo.styles['List Bullet']
                sec3a.text = (memo_text)
                sec3a.add_run().add_break()
                sec3a.add_run().add_break()
                sec3a.add_run().add_text('Recommendation: ')
                sec3a.add_run().add_break()
                sec3b = docMemo.tables[2].cell(1,0).paragraphs[1].add_run(reco_text)
                sec3a.add_run().add_break()
                sec3b.bold = False
                sec3a.insert_paragraph_before(text=None, style=None)
                docMemo.tables[1].cell(2,2).add_paragraph().add_run().add_break()
                docMemo.tables[1].cell(2,2).add_paragraph().add_run().add_text('See below for more details of the issues noted.')

            ###setting variable paths
            #set screenshot path
            if not os.path.exists(output_path + '\screenshots'):
                os.makedirs(output_path + '\screenshots')
            screenshot_path = r'\screenshots'
            
            ###open the delivery Memo
            TPA = doc_fr['TPAName'][row_num].replace('_', '').replace('apac','')
            templates = os.listdir(r'Templates')
            templates = [x.lower() for x in templates]
            template = [temp for temp in templates if TPA.lower() in temp.replace(' ','')][0]
            os.chdir(r'Templates')
            docMemo = docx.Document(template)

            #Table 1 column widths setup
            for cell in docMemo.tables[1].columns[0].cells:
                cell.width = Inches(2.44)

            for cell in docMemo.tables[1].columns[1].cells:
                cell.width = Inches(1)

            for cell in docMemo.tables[1].columns[2].cells:
                cell.width = Inches(3.04)


            #fill up engagement name
            docMemo.tables[0].cell(0,1).text ='%s (%s)'%(doc_fr['FundName'][row_num], doc_fr['FundShortCode'][row_num])

            #fill up year-end
            try: 
                d = datetime.strptime(doc_fr['FinancialEndDate'][row_num], '%d-%m-%y')
                d_start = datetime.strptime(doc_fr['FinancialStartDate'][row_num], '%d-%m-%y')
            except ValueError:
                d = datetime.strptime(doc_fr['FinancialEndDate'][row_num], '%d/%m/%Y')
                d_start = datetime.strptime(doc_fr['FinancialStartDate'][row_num], '%d/%m/%Y')

            fileDate = d.strftime('%m%Y')
            
            #for period in filename of delivery memo
            period_days = d - d_start
            period = round(period_days.days/30)
            if period < 10:
                period = '0'+str(period)
            else:
                period = str(period)

            docMemo.tables[0].cell(2,1).text= d.strftime('%d %B %Y')
            #fill up date of documentation
            currentDate = datetime.now()
            docMemo.tables[0].cell(4,1).text =currentDate.strftime('%d %B %Y')
            #docMemo.tables[0].cell(6,1).text =currentDate.strftime('%d %B %Y')
            #fill up preparer & reviewer
            docMemo.tables[0].cell(3,1).text = Preparer
            docMemo.tables[0].cell(5,1).text = Reviewer


            style = docMemo.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(10)
            font.color.rgb = RGBColor(0, 0, 0)

            ###For leads
            account_names_append = []
            account_class_append = []

            ### set variables of the WAMapps source files to be used
            #set file
            wp_path = '/Workpapers'
            files = os.listdir(output_path+'/'+wp_path)
            os.chdir(output_path+'/'+wp_path)

            WAMapps_Leads = [file for file in files if "%s_%s_Lead_Sheets" %(str(doc_fr['FundShortCode'][row_num])[0:10],fileDate) in file]
            WAMapps_CapAct = [file for file in files if "%s_%s_Capital_Activity" %(str(doc_fr['FundShortCode'][row_num])[0:10],fileDate) in file]
            WAMapps_CashRecon = [file for file in files if "%s_%s_Cash_Reconciliation" %(str(doc_fr['FundShortCode'][row_num])[0:10],fileDate) in file]
            WAMapps_Dividends = [file for file in files if "%s_%s_Dividends" %(str(doc_fr['FundShortCode'][row_num])[0:10],fileDate) in file]
            WAMapps_Interest = [file for file in files if "%s_%s_Interest" %(str(doc_fr['FundShortCode'][row_num])[0:10],fileDate) in file]
            WAMapps_InvRF = [file for file in files if "%s_%s_Investment_Roll" %(str(doc_fr['FundShortCode'][row_num])[0:10],fileDate) in file]
            WAMapps_PortVal = [file for file in files if "%s_%s_Portfolio_Val" %(str(doc_fr['FundShortCode'][row_num])[0:10],fileDate) in file]
            WAMapps_RGL = [file for file in files if "%s_%s_Realized_Gain" %(str(doc_fr['FundShortCode'][row_num])[0:10],fileDate) in file]
            WAMapps_Trades = [file for file in files if "%s_%s_Unsettled_Trades" %(str(doc_fr['FundShortCode'][row_num])[0:10],fileDate) in file]
            WAMapps_FXReason = [file for file in files if "%s_%s_FX_Reasonability" %(str(doc_fr['FundShortCode'][row_num])[0:10],fileDate) in file]
            WAMapps_NAVbased = [file for file in files if "%s_%s_NAV_Based" %(str(doc_fr['FundShortCode'][row_num])[0:10],fileDate) in file]
            
            #validation - if fund code and end date did not match any workpapers
            if WAMapps_Leads == WAMapps_CapAct == WAMapps_CashRecon == WAMapps_Dividends == WAMapps_Interest == WAMapps_InvRF == \
            WAMapps_PortVal == WAMapps_RGL == WAMapps_Trades == WAMapps_FXReason == WAMapps_NAVbased == []:
                print("\n(FR Row Number " + str(row_num) + ") Invalid Run: " + "\nThe specified fund code and financial end date in the FR did not match any of the workpapers in the filepath.")
                print('Workpaper filename the tool is looking for: '+ str(doc_fr['ClientName'][row_num])[0:10] + '_' + str(doc_fr['FundShortCode'][row_num])[0:10] + '_' + fileDate + '_Workpaper.xlsx')
                print('Make sure that the client name (before the first underscore of filename) indicated in the workpapers is: ' + '\"' + str(doc_fr['ClientName'][row_num])[0:10] + '\".')
                print('Make sure that the fund code (after the first underscore of filename) indicated in the workpapers is: ' + '\"' + str(doc_fr['FundShortCode'][row_num])[0:10] + '\".')
                print('Make sure that the period end date (after the second underscore of filename) indicated in the workpapers is: ' + '\"' + fileDate + '\".')
                
                # #Delete all current variables to avoid bugs when batch processing different requests
                # for var in dir():
                #     if var not in Global_var:
                #         del globals()[var]
                        
            else:
                #To Remove Bug of Not Running
                pyautogui.FAILSAFE = False
                pyautogui.click(3000,3000)
                time.sleep(.5)

                import pythoncom
                try:
                    xl = client.gencache.EnsureDispatch('Excel.Application',pythoncom.CoInitialize())
                    # xl = win32com.client.Dispatch("Excel.Application",pythoncom.CoInitialize())
                except AttributeError:
                    # Corner case dependencies.
                    import os
                    import re
                    import sys
                    import shutil
                    # Remove cache and try again.
                    MODULE_LIST = [m.__name__ for m in sys.modules.values()]
                    for module in MODULE_LIST:
                        if re.match(r'win32com\.gen_py\..+', module):
                            del sys.modules[module]
                    shutil.rmtree(os.path.join(os.environ.get('LOCALAPPDATA'), 'Temp', 'gen_py'))
                    from win32com import client
                    xl = client.gencache.EnsureDispatch('Excel.Application',pythoncom.CoInitialize())

                os.chdir(output_path)
                os.chdir(r'Templates')

                doc = docx.Document('FY20 Transformation Tracker.docx')

                #fill up engagement name
                doc.tables[0].cell(0,1).text= '%s (%s)'%(doc_fr['FundName'][row_num], doc_fr['FundShortCode'][row_num])
                doc.tables[0].cell(1,1).text= d.strftime('%d %B %Y')
                
                #fill up date of documentation
                currentDate = datetime.now()
                doc.tables[0].cell(2,1).text = Preparer
                doc.tables[0].cell(3,1).text =currentDate.strftime('%d %B %Y')
                doc.tables[0].cell(4,1).text = Reviewer

                style = doc.styles['Normal']
                font = style.font
                font.name = 'Calibri'
                font.size = Pt(10)

                #convert date to string to be used in filename
                dateString = str(d.year)[2:]

                ##save the transformation tracker memo
                os.chdir(output_path)
                doc.save('FY%s %s (%s) Transformation.docx' %(dateString,doc_fr['FundName'][row_num], doc_fr['FundShortCode'][row_num]))
                transformation_file = 'FY%s %s (%s) Transformation.docx' %(dateString,doc_fr['FundName'][row_num], doc_fr['FundShortCode'][row_num])

                ################LEADS WORKPAPER TESTING#######################################################
                if len(WAMapps_Leads)!=0:
                    xl=client.gencache.EnsureDispatch("Excel.Application")
                    xl.Visible = True
                    xl.WindowState = client.constants.xlMaximized
                    xl.FormulaBarHeight = 2
                    #xl.ActiveWindow.Zoom = 70
                    workbook = xl.Workbooks.Open(output_path+'/'+wp_path+'/'+WAMapps_Leads[0])
                    workbook.Save()

                    #trial balance CY & PY test
                    try:
                        Leads_TB = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_Leads[0], sheet_name='A_Trial_Balance', header=header_no, skipfooter=1)
                        Leads_TBPivot = pd.pivot_table(Leads_TB[Leads_TB['Share Class'] == 'Fund Level'], index=['Tab'], aggfunc='sum', values=['Current_Year', 'Previous_Year'])
                        Leads_TBClass = pd.pivot_table(Leads_TB[Leads_TB['Share Class'] == 'Fund Level'], index=['Class', 'Tab'], aggfunc='sum', values=['Current_Year', 'Previous_Year'])
                        Leads_Sheets = pd.ExcelFile(output_path+'/'+wp_path+'/'+WAMapps_Leads[0]) 
                    except:
                        TB_NotFound = True

                    list_tabs = []
                    for tab in ['C1','D1','G1','N1','T1','UB1','VD1','Z1','Z2']:
                        if tab in Leads_Sheets.sheet_names:
                            list_tabs.append(tab)
                        else:
                            pass

                    #Screenshots of each tab for recommendation
                    for tab in list_tabs:
                        if tab in Leads_TBPivot.index.tolist():
                            accountnames = Leads_TB[Leads_TB['Tab']==tab]['Account_Name'].tolist()
                            accountclass = set(Leads_TB[Leads_TB['Tab']==tab]['Class'].tolist())
                            os.chdir(output_path+'/'+screenshot_path)
                            workbook.Sheets(tab).Select()
                            workbook.Sheets(tab).Cells(1,5).ColumnWidth = 17.22
                            workbook.Sheets(tab).Cells(1,7).ColumnWidth = 17.22
                            workbook.Sheets(tab).Cells(1,9).ColumnWidth = 19.22
                            width = workbook.Sheets(tab).Cells(1,3).ColumnWidth
                            for row_num2 in range(workbook.Sheets(tab).UsedRange.Rows.Count):
                                col_num = 2 # Fixed to look at column 2 only
                                # Note Python's range() counts from 0 and Excel counts from 1
                                value = workbook.Sheets(tab).Cells(row_num2 + 1, col_num + 1).Value
                                if value in accountclass:
                                    workbook.Sheets(tab).Cells(row_num2+1, col_num+1).Select()
                                    pyautogui.hotkey('ctrl', 'down')
                                    pyautogui.hotkey('ctrl', 'down')
                                    time.sleep(0.5)
                                    class1 = pyautogui.screenshot(re.sub('[^A-Za-z0-9]+', '', value)+'.png',region=(26,430, width*12.65+380, 540))

                    class2 = [class1[:-4] for class1 in os.listdir(output_path+'/'+screenshot_path)]

                    if (abs(Leads_TBPivot.sum())>1).any():
                        os.chdir(output_path+'/'+screenshot_path)
                        #navigate thru the workbook
                        workbook.Sheets('A_Trial_Balance').Select()
                        workbook.Sheets('A_Trial_Balance').Cells(1,1).Select()
                        #in case it is in middle of woorkbook
                        workbook.Sheets('A_Trial_Balance').Cells(11,2).Select()
                        xl.ActiveWindow.Zoom = 70
                        pyautogui.hotkey('ctrl','down')
                        pyautogui.press('down',presses=11)
                        workbook.Sheets('A_Trial_Balance').Cells(1,2).ColumnWidth = 14
                        workbook.Sheets('A_Trial_Balance').Cells(1,3).ColumnWidth = 13
                        workbook.Sheets('A_Trial_Balance').Cells(1,4).ColumnWidth = 14
                        workbook.Sheets('A_Trial_Balance').Cells(1,5).ColumnWidth = 40
                        workbook.Sheets('A_Trial_Balance').Cells(1,6).ColumnWidth = 16
                        workbook.Sheets('A_Trial_Balance').Cells(1,7).ColumnWidth = 4
                        workbook.Sheets('A_Trial_Balance').Cells(1,8).ColumnWidth = 30
                        workbook.Sheets('A_Trial_Balance').Cells(1,11).ColumnWidth = 18.5
                        workbook.Sheets('A_Trial_Balance').Cells(1,12).ColumnWidth = 18.5
                        time.sleep(1)
                        TBnotzero = pyautogui.screenshot('TBnotzero1.png',region=(87,462,1285,315)) #Edited region=(50,300, 1750, 660))
                        TBnotzero = add_border(TBnotzero,'TBnotzero1.png')

                        if (Leads_TBPivot['Current_Year'].sum())>1 and (Leads_TBPivot['Previous_Year'].sum())>1:
                            sec3 = 'The current & prior year trial balances does not reconcile.'
                            leads_reco = 'The audit team shall investigate the prior & current year trial balance input files and determine the root cause of the issue.'
                        elif (Leads_TBPivot['Previous_Year'].sum())>1:
                            sec3 = 'The prior year trial balance does not reconcile.'
                            leads_reco = 'The audit team shall investigate the prior year trial balance input file and determine the root cause of the issue.'
                        else:
                            sec3 = 'The current year trial balance does not reconcile.'
                            leads_reco = 'The audit team shall investigate the current year trial balance input file and determine the root cause of the issue.'

                        memo_final(memo_text = sec3,screenshot = [('TBnotzero1.png',Inches(6))],reco_text=leads_reco)

                    #trial balance D1 test
                    try:
                        Leads_D1 = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_Leads[0], sheet_name='D1', header=header_no, skipfooter=1)
                        Leads_tieOut = Leads_D1.tail(2)
                        Leads_z2tieOut_sum = Leads_tieOut['CY Cost Value'].sum()
                        Leads_TB['D1Diff'] = Leads_TB['Current_Year'].apply(lambda x: np.nan if x == 0 else abs(x)- abs(Leads_z2tieOut_sum))
                        Leads_TBClass['D1ClassDiff'] = Leads_TBClass['Current_Year'].apply(lambda x: np.nan if x == 0 else abs(x)- abs(Leads_z2tieOut_sum))
                        Leads_TB['D1_Reclass'] = Leads_TB['D1Diff'].apply(lambda x: 'Reclass' if abs(x)<1 else 'No Reclass')
                        Leads_TBClass['D1_Reclass'] = Leads_TBClass['D1ClassDiff'].apply(lambda x: 'Reclass' if abs(x)<1 else 'No Reclass')

                        D1reco = []
                        reco_ss = []
                        d1_reclass_tab = []

                        if abs(Leads_z2tieOut_sum)>1:
                            if (Leads_D1['CY Unrealised Value'].fillna(0)==0).all() and (Leads_D1['PY Unrealised Value'].fillna(0)==0).all():
                                D1reco.append('The difference is because no D1 account is mapped as UNREALIZED value. As such, audit team may opt to remap the said account classes.')
                            elif (Leads_TB['D1_Reclass']=='Reclass').any():
                                account_name = Leads_TB[Leads_TB['D1_Reclass']=='Reclass']['Account_Name'].iloc[0]
                                account_class = Leads_TB[Leads_TB['D1_Reclass']=='Reclass']['Class'].iloc[0]
                                d1_reclass_tab = Leads_TB[Leads_TB['D1_Reclass']=='Reclass']['Tab'].iloc[0]
                                D1reco.append('The difference in D1 is due to account: %s. As such the audit team may opt to manually reclassify the said account from %s to its appropriate account class.'%(account_name, account_class))
                                account_names_append.append(account_name)
                                account_class_append.append(account_class)
                                if re.sub('[^A-Za-z0-9]+', '',account_class) in class2:
                                    reco_ss.append(re.sub('[^A-Za-z0-9]+', '',account_class)+'.png')
                                else:
                                    pass
                            elif (Leads_TBClass['D1_Reclass']=='Reclass').any():
                                account_class = Leads_TBClass[Leads_TBClass['D1_Reclass']=='Reclass'].index.values.tolist()[0][0]
                                d1_reclass_tab = Leads_TBClass[Leads_TBClass['D1_Reclass']=='Reclass'].index.values.tolist()[0][0]
                                D1reco.append('The difference in D1 is due to account class: %s. As such the audit team may opt to manually reclassify the said account class.'%(account_class))
                                account_class_append.append(account_class)
                                if re.sub('[^A-Za-z0-9]+', '',account_class) in class2:
                                    reco_ss.append(re.sub('[^A-Za-z0-9]+', '',account_class)+'.png')
                                else:
                                    pass
                            else:
                                D1reco.append('The difference is probably due to topside adjustments that are not yet accounted for. As such, the audit team shall manually reconcile the said difference once the topside adjustments become available.') 

                            D1reco_str = ' '.join(D1reco)

                            os.chdir(output_path+'/'+screenshot_path)
                            workbook.Sheets('D1').Select()
                            xl.ActiveWindow.Zoom = 70
                            lastrow = workbook.Sheets('D1').UsedRange.Rows.Count
                            workbook.Sheets('D1').Cells(lastrow,1).Select()
                            pyautogui.hotkey('ctrl', 'down')
                            pyautogui.press('right')
                            pyautogui.hotkey('ctrl', 'up')
                            pyautogui.hotkey('ctrl', 'up')
                            workbook.Sheets('D1').Cells(1,2).ColumnWidth = 18
                            workbook.Sheets('D1').Cells(1,3).ColumnWidth = 14
                            workbook.Sheets('D1').Cells(1,4).ColumnWidth = 18.5
                            workbook.Sheets('D1').Cells(1,6).ColumnWidth = 5
                            time.sleep(1)
                            D1tieOutnotzero = pyautogui.screenshot('D1tieOutnotzero.png',region=(55, 582, 490, 125)) #Edited
                            D1tieOutnotzero = add_border(D1tieOutnotzero,'D1tieOutnotzero.png')

                            memo_final(memo_text = 'The "Tie out to Z2 Lead" in D1 tab of the Leads working paper does not reconcile.',screenshot = [('D1tieOutnotzero.png',Inches(3.5))],reco_text=D1reco_str)
                            lead_inputs = [(reco_ss,d1_reclass_tab,'D1')]
                            refer_memo(lead_inputs)
                    except:
                        D1_NotFound = True

                    #trial balance T1 test
                    try:
                        Leads_T1 = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_Leads[0], sheet_name='T1', header=header_no, skipfooter=1)
                        T1_plBal= Leads_T1.tail(1)
                        T1_rectoBS = Leads_T1.tail(7).head(1)
                        T1_recon = T1_plBal['CY Balance per TB'].iloc[0]+T1_rectoBS['CY Balance per TB'].iloc[0]
                    except:
                        T1_NotFound = True

                    if abs(T1_recon)>1:
                        os.chdir(output_path+'/'+screenshot_path)
                        workbook.Sheets('T1').Select()
                        xl.ActiveWindow.Zoom = 70
                        lastrow = workbook.Sheets('T1').UsedRange.Rows.Count
                        workbook.Sheets('T1').Cells(lastrow,1).Select()
                
                        pyautogui.hotkey('ctrl', 'down')
                        pyautogui.press('right',presses=2)
                        pyautogui.hotkey('ctrl', 'up')
                        pyautogui.hotkey('ctrl', 'up')
                        pyautogui.hotkey('ctrl', 'up')
                        pyautogui.hotkey('ctrl', 'up')
                        pyautogui.hotkey('ctrl', 'up')
                        workbook.Sheets('T1').Cells(1,3).ColumnWidth = 36
                        time.sleep(1)
                        T1tieOutnotzero = pyautogui.screenshot('T1tieOutnotzero.png',region=(167, 606, 435, 295))# Edited region=(130,450, 1000, 520))
                        T1tieOutnotzero = add_border(T1tieOutnotzero,'T1tieOutnotzero.png')

                        memo_final(memo_text = 'Reconciliation to Balance Sheet does not reconcile with Profit/Loss from operations in T1 tab.',screenshot = [('T1tieOutnotzero.png',Inches(3))],reco_text='The audit team shall investigate the trial balance input file or the account mapping of some accounts as this may be the reason of the variance.')

                    ###Share Capital Rec Test 
                    try:
                        Leads_ShareCap = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_Leads[0], sheet_name='Share_Capital_Rec', header=header_no)
                    except:
                        ShareCap_NotFound = True

                    try:
                        if (abs(Leads_ShareCap['Difference in Shares'])>1).any():
                            workbook.Sheets('Share_Capital_Rec').Select()
                            xl.ActiveWindow.Zoom = 70
                            #For row height (region length)
                            ShareRec_Filter = 0
                            ShareRec_Filter = sum(abs(Leads_ShareCap['Difference in Shares'])>1)+1
                            #Limiting screenshot length if too much variance
                            if ShareRec_Filter >= 21:
                                ShareRec_height = 21 * 18
                            elif ShareRec_Filter <= 7:
                                ShareRec_height = ShareRec_Filter * 22
                            else:
                                ShareRec_height = ShareRec_Filter * 18
                            workbook.Sheets('Share_Capital_Rec').Cells(1,3).ColumnWidth = 40
                            workbook.Sheets('Share_Capital_Rec').Cells(header_no+1,2).Select()
                            pyautogui.hotkey('ctrl', 'shift','end')
                            pyautogui.hotkey('alt', 'a','t')
                            pyautogui.hotkey('ctrl', 'end','ctrl', 'up','ctrl', 'up','ctrl', 'down')
                            pyautogui.hotkey('alt', 'down')
                            pyautogui.hotkey('f', 'g','1','tab','right','tab')
                            pyautogui.press('down',presses=6)
                            pyautogui.hotkey('enter','tab','-','1','enter')
                            pyautogui.press('left',presses=4)
                            pyautogui.hotkey('shift','left','shift','left')
                            pyautogui.hotkey('ctrl','0','ctrl','left','up','up','left')
                            time.sleep(1)

                            ShareCapRecNotzero = pyautogui.screenshot('ShareCapRecNotzero.png',region=(49,457,1420,ShareRec_height))#Edited region=(50,300, 1800, 320))
                            ShareCapRecNotzero = add_border(ShareCapRecNotzero,'ShareCapRecNotzero.png')

                            memo_final(memo_text = 'The shares outstanding at year end as per input file do not reconcile with the calculated outstanding year end shares.',screenshot = [('ShareCapRecNotzero.png',Inches(6))],reco_text='The audit team shall investigate the share register & capital activity input files and determine the reason of the variance.')
                    except:
                        pass

                    ###Account Rec & Balance Sheet FS testing 
                    try:
                        Leads_AccountRec = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_Leads[0], sheet_name='Account_Rec', header=header_no)
                        Leads_BS_FS = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_Leads[0], sheet_name='Balance_Sheet_FS', header=header_no)
                        TBBal_total = abs(Leads_AccountRec[Leads_AccountRec['Account Number']=='Total']['TB Balance'].iloc[0].round(2))
                        FSBal_total = abs(Leads_AccountRec[Leads_AccountRec['Account Number']=='Total']['FS Balance'].iloc[0].round(2))
                        BSFS_total = abs(Leads_BS_FS[Leads_BS_FS['FS Account Class']=='Total']['CY Balance'].iloc[0].round(2))
                        AdjBSFS_total = abs(Leads_BS_FS[Leads_BS_FS['FS Account Class']=='Total']['Adjusted FS Balance'].iloc[0].round(2))
                    except:
                        AccountRec_NotFound = True
                        BS_FS_NotFound = True

                    try:
                        if (TBBal_total>1 or FSBal_total>1) or (BSFS_total>1 or AdjBSFS_total>1):
                            #to screenshot in account rec tab
                            os.chdir(output_path+'/'+screenshot_path)
                            workbook.Sheets('Account_Rec').Select()
                            xl.ActiveWindow.Zoom = 70
                            workbook.Sheets('Account_Rec').Cells(1,2).ColumnWidth = 18.22
                            workbook.Sheets('Account_Rec').Cells(1,3).ColumnWidth = 39.22
                            workbook.Sheets('Account_Rec').Cells(1,4).ColumnWidth = 57.22
                            workbook.Sheets('Account_Rec').Cells(1,5).ColumnWidth = 57.22
                            workbook.Sheets('Account_Rec').Cells(1,6).ColumnWidth = 18.22
                            workbook.Sheets('Account_Rec').Cells(1,7).ColumnWidth = 4.22
                            workbook.Sheets('Account_Rec').Cells(1,8).ColumnWidth = 18.22
                            lastrow = workbook.Sheets('Account_Rec').UsedRange.Rows.Count
                            workbook.Sheets('Account_Rec').Cells(lastrow,1).Select()
                            time.sleep(1)
                            AccountRecnotZero = pyautogui.screenshot('AccountRecnotZero.png',region=(50,430, 1800, 380))
                            AccountRecnotZero = add_border(AccountRecnotZero,'AccountRecnotZero.png')

                            #to screenshot in bsfs tab
                            workbook.Sheets('Balance_Sheet_FS').Select()
                            xl.ActiveWindow.Zoom = 70
                            workbook.Sheets('Balance_Sheet_FS').Cells(1,2).ColumnWidth = 75.22
                            workbook.Sheets('Balance_Sheet_FS').Cells(1,3).ColumnWidth = 34.22
                            workbook.Sheets('Balance_Sheet_FS').Cells(1,4).ColumnWidth = 17.22
                            workbook.Sheets('Balance_Sheet_FS').Cells(1,5).ColumnWidth = 14.22
                            workbook.Sheets('Balance_Sheet_FS').Cells(1,6).ColumnWidth = 17.22
                            workbook.Sheets('Balance_Sheet_FS').Cells(1,7).ColumnWidth = 4.22
                            workbook.Sheets('Balance_Sheet_FS').Cells(1,8).ColumnWidth = 17.22
                            lastrow = workbook.Sheets('Balance_Sheet_FS').UsedRange.Rows.Count
                            workbook.Sheets('Balance_Sheet_FS').Cells(lastrow,1).Select()
                            time.sleep(1)
                            BSFSnotZero = pyautogui.screenshot('BSFSnotZero.png',region=(50,430, 1800, 480))
                            BSFSnotZero = add_border(BSFSnotZero,'BSFSnotZero.png')

                            if (TBBal_total>1 or FSBal_total>1) and (BSFS_total>1 or AdjBSFS_total>1):
                                sec3 = 'The totals in the Account_Rec and Balance_Sheet_FS tabs of the Leads working paper do not net to nil.'
                                accbal_ss = [('AccountRecnotZero.png',Inches(5.8)),('BSFSnotZero.png',Inches(5.8))]
                            elif (TBBal_total>1 or FSBal_total>1) and (BSFS_total<=1 and AdjBSFS_total<=1):
                                sec3 = 'The totals in the Account_Rec tab of the Leads working paper does not net to nil.'
                                accbal_ss = [('AccountRecnotZero.png',Inches(5.8))]
                            elif (TBBal_total<=1 and FSBal_total<=1) and (BSFS_total>1 or AdjBSFS_total>1):
                                sec3 = 'The totals in the Balance_Sheet_FS tab of the Leads working paper does not net to nil.'
                                accbal_ss = [('BSFSnotZero.png',Inches(5.8))] 
                            else:
                                pass

                            memo_final(memo_text = sec3,screenshot = accbal_ss,reco_text='The audit team shall investigate the trial balance and adjustments input files and determine the root cause of the variance.')

                    except:
                        pass

                    ###Statement of Changes Test 
                    try:
                        Leads_SOC = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_Leads[0], sheet_name='Statement_of_Changes', header=header_no, skipfooter=1)
                        SOC_basedonNA = Leads_SOC[Leads_SOC['Unnamed: 1']=='Increase/decrease in net assets']['Raw'].iloc[0]
                        SOC_basedonNAV = Leads_SOC[Leads_SOC['Unnamed: 1']=='Increase/decrease in fund assets based on NAV']['Raw'].iloc[0]
                        SOC_recon = (SOC_basedonNA-SOC_basedonNAV).round(2)
                    except:
                        SOC_NotFound = True

                    try:
                        if abs(SOC_recon)>1:
                            os.chdir(output_path+'/'+screenshot_path)
                            workbook.Sheets('Statement_of_Changes').Select()
                            xl.ActiveWindow.Zoom = 70
                            lastrow = workbook.Sheets('Statement_of_Changes').UsedRange.Rows.Count
                            workbook.Sheets('Statement_of_Changes').Cells(lastrow,1).Select()
                            pyautogui.hotkey('ctrl', 'down')
                            pyautogui.press('right')
                            pyautogui.hotkey('ctrl','up')
                            pyautogui.press('up',presses=11)
                            workbook.Sheets('Statement_of_Changes').Cells(1,2).ColumnWidth = 68
                            workbook.Sheets('Statement_of_Changes').Cells(1,3).ColumnWidth = 20
                            time.sleep(1)
                            SOCnotZero = pyautogui.screenshot('SOCnotZero.png',region=(54, 458,630, 240))#Edited region=(50,300, 920, 660))
                            SOCnotZero = add_border(SOCnotZero,'SOCnotZero.png')

                            memo_final(memo_text = 'There is a difference noted in the "Statement of Changes" in the Leads working paper.',screenshot = [('SOCnotZero.png',Inches(4.5))],reco_text='The audit team shall investigate the trial balance and capital activity input files and determine the root cause of the variance.')
                    except:
                        pass

                    workbook.Close(SaveChanges=False)


                ################CAPITAL ACTIVITY WORKPAPER TESTING#######################################################

                if len(WAMapps_CapAct)!=0:
                    xl=client.gencache.EnsureDispatch("Excel.Application")
                    xl.Visible = True
                    xl.WindowState = client.constants.xlMaximized
                    xl.FormulaBarHeight = 2
                    workbook = xl.Workbooks.Open(output_path+'/'+wp_path+'/'+WAMapps_CapAct[0])
                    workbook.Save()

                    try:
                        CapAct_Summary = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_CapAct[0], sheet_name='Summary', header=header_no, skipfooter=1)
                        CapAct_temp0=CapAct_Summary[CapAct_Summary['Accrual Testing Key (Click + to expand)']=='Difference'].fillna(0)
                        CapAct_Diff1=CapAct_temp0['Unnamed: 2']
                        #compute difference of last row
                        CapAct_temp1 = CapAct_Summary.tail(2)
                        CapAct_temp2 = CapAct_temp1.head(1).fillna(0)
                        CapAct_temp3 = CapAct_temp1.tail(1).fillna(0)
                        CapAct_Diff2=CapAct_temp2['Unnamed: 2'].iloc[0]+CapAct_temp3['Unnamed: 2'].iloc[0]
                        #for performance of step2
                        SC_T1 = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_CapAct[0], sheet_name='Share_Capital_T1', header=header_no+3)
                        CA_Detail = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_CapAct[0], sheet_name='Capital Activity', header=header_no)
                        sc_transactionClass = CA_Detail.groupby(['Transaction Class']).sum()['Net Value (Base)']
                    except:
                        CapAct_NotFound = True

                    if (abs(CapAct_Diff1)>1).any() or abs(CapAct_Diff2)>1:
                        os.chdir(output_path+'/'+screenshot_path)
                        #navigate thru the workbook
                        workbook.Sheets('Summary').Select()
                        xl.ActiveWindow.Zoom = 60
                        xl.ActiveWindow.FreezePanes = False
                        workbook.Sheets('Summary').Cells(1,1).Select()
                
                        pyautogui.press('down', presses=60)
                        workbook.Sheets('Summary').Cells(1,3).ColumnWidth = 22
                        time.sleep(1)
                        CapActSumNotZero = pyautogui.screenshot('CapActSumNotZero.png',region=(80,305, 430, 665))
                        CapActSumNotZero = add_border(CapActSumNotZero,'CapActSumNotZero.png')

                        #to define the differences per mapping
                        ca_names = []
                        ca_reco = []
                        reco_subs = []
                        reco_reds = []
                        reco_transfers = []
                        reco_dist = []
                        reco_other = []
                        subs_reclass_tab = [] 
                        reds_reclass_tab = [] 
                        transfers_reclass_tab = [] 
                        dist_reclass_tab = [] 
                        other_reclass_tab = []
                        ca_inputs = [(CapAct_Diff1.iloc[0],ca_names,'Subscriptions','Subscriptions','Subscriptions_Reclass',Leads_Sheets,'Coversheet',ca_reco,subs_reclass_tab,reco_subs)
                                    ,(CapAct_Diff1.iloc[1],ca_names,'Redemptions','Redemptions','Redemptions_Reclass',Leads_Sheets,'Coversheet',ca_reco,reds_reclass_tab,reco_reds)
                                    ,(CapAct_Diff1.iloc[2],ca_names,'Transfers','Transfers','Transfers_Reclass',Leads_Sheets,'Coversheet',ca_reco,transfers_reclass_tab,reco_transfers)
                                    ,(CapAct_Diff1.iloc[3],ca_names,'Distributions','Distributions','Dist_Reclass',Leads_Sheets,'Coversheet',ca_reco,dist_reclass_tab,reco_dist)
                                    ,(CapAct_Diff2,ca_names,'Other Capital','Other Capital','Other Capital',Leads_Sheets,'Coversheet',ca_reco,other_reclass_tab,reco_other)]

                        recommendation(ca_inputs)

                        ##for proper preposition and grammar
                        if len(ca_names)>1:
                            preposition = 'do not reconcile with their Leads counterparts'
                        else:
                            preposition = 'does not reconcile with its Leads counterpart'

                        ca_names_str = ', '.join(ca_names)
                        ca_reco_str = ' \n\n'.join(ca_reco)

                        memo_final(memo_text ='%s per capital activity working paper %s.'%(ca_names_str,preposition),screenshot = [('CapActSumNotZero.png',Inches(2.5))],reco_text=ca_reco_str)
                        ca_inputs = [(reco_subs,subs_reclass_tab,'subscriptions'),(reco_reds,reds_reclass_tab,'redemptions'),(reco_transfers,transfers_reclass_tab,'transfers'),(reco_dist,dist_reclass_tab,'distributions'),(reco_other,other_reclass_tab,'other capital')]
                        refer_memo(ca_inputs)


                    #Share Capital T1 testing
                    try:
                        CapAct_T1 = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_CapAct[0], sheet_name='Share_Capital_T1', header=header_no, skipfooter=1)
                        NAVperTA = CapAct_T1[CapAct_T1['Unnamed: 3']=='NAV per TA']['Unnamed: 5'].fillna(0).iloc[0]
                        NAVperTB = CapAct_T1[CapAct_T1['Unnamed: 3']=='NAV per TB']['Unnamed: 5'].fillna(0).iloc[0]
                        PLperTA = CapAct_T1[CapAct_T1['Unnamed: 3']=='(Profit)/Loss per TA']['Unnamed: 5'].fillna(0).iloc[0]
                        PLperTB = CapAct_T1[CapAct_T1['Unnamed: 3']=='(Profit)/Loss per TB']['Unnamed: 5'].fillna(0).iloc[0]
                        try:
                            PYunrTB = CapAct_T1[CapAct_T1['Unnamed: 3']=='PY Unr per TB']['Unnamed: 5'].fillna(0).iloc[0]
                        except:
                            PYunrTB = 0
                        NAVdiff = (NAVperTA-NAVperTB).round(2)
                        PLdiff = (PLperTA+PLperTB-PYunrTB).round(2)
                    except:
                        CapAct_T1NotFound = True

                    if abs(NAVdiff)>1 or abs(PLdiff)>1:
                        os.chdir(output_path+'/'+screenshot_path)
                        workbook.Sheets('Share_Capital_T1').Select()
                        xl.ActiveWindow.Zoom = 70
                        xl.ActiveWindow.FreezePanes = False
                        workbook.Sheets('Share_Capital_T1').Cells(1,1).Select()
                        workbook.Sheets('Share_Capital_T1').Cells(1,3).ColumnWidth = 24.33
                        #in case it is in middle of woorkbook
                        # pyautogui.press('down',presses=54)
                        pyautogui.hotkey('ctrl', 'down')
                        pyautogui.hotkey('right')
                        pyautogui.hotkey('ctrl', 'up')
                        time.sleep(1)
                        T1Diff = pyautogui.screenshot('T1Diff.png',region=(425,365,515,260))#Edited region=(400,450, 600, 520))
                        T1Diff = add_border(T1Diff,'T1Diff.png')
                        if abs(NAVdiff)>1 and abs(PLdiff)>1:
                            T1_reason = 'NAV per TA does not reconcile with NAV per TB and Profit/Loss per TA does not reconcile with Profit/Loss per TB.'
                        elif abs(NAVdiff)>1:
                            T1_reason = 'NAV per TA does not reconcile with NAV per TB.'
                        else:
                            T1_reason = 'Profit/Loss per TA does not reconcile with Profit/Loss per TB.'

                        memo_final(memo_text =T1_reason,screenshot = [('T1Diff.png',Inches(3.5))],reco_text='The audit team shall investigate the capital activity files and the mapping of the trial balance accounts as these are probably reasons of the variance.')

                    #Reconciliation to T1 Testing
                    try:
                        quantityDiff = abs(CapAct_T1[CapAct_T1['ALLOCATION SHEETS: REC TO T1']=='Difference']['Unnamed: 3'].iloc[0])
                        valueDiff = abs(CapAct_T1[CapAct_T1['ALLOCATION SHEETS: REC TO T1']=='Difference']['Unnamed: 7'].iloc[0])
                        series1=pd.Series(CapAct_T1['Unnamed: 3'])
                        T1quantity_col = pd.to_numeric(series1, errors='coerce').fillna(0)
                    except:
                        CapAct_T1NotFound = True

                    if abs(quantityDiff)>1 or abs(valueDiff)>1:
                        os.chdir(output_path+'/'+screenshot_path)
                        workbook.Sheets('Share_Capital_T1').Select()
                        xl.ActiveWindow.Zoom = 70
                        xl.ActiveWindow.FreezePanes = False
                        workbook.Sheets('Share_Capital_T1').Cells(1,3).ColumnWidth = 24.33
                        workbook.Sheets('Share_Capital_T1').Cells(1,1).Select()
                        #in case it is in middle of woorkbook
                        time.sleep(1)
                        sharecapT1Diff = pyautogui.screenshot('sharecapT1Diff.png',region=(26,465, 1250, 485))
                        sharecapT1Diff = add_border(sharecapT1Diff,'sharecapT1Diff.png')
                        if abs(quantityDiff)>1 and abs(valueDiff)>1:
                            T1rec_reason = 'Quantity & Amount reconciliation in Share_Capital_T1 tab of the capital activity workpaper do not net to nil.'
                        elif abs(quantityDiff)>1 and abs(valueDiff)==0:
                            T1rec_reason = 'Quantity reconciliation in Share_Capital_T1 tab of the capital activity workpaper does not net to nil.'
                        else:
                            T1rec_reason = 'Amount reconciliation in Share_Capital_T1 tab of the capital activity workpaper does not net to nil.'

                        memo_final(memo_text =T1rec_reason,screenshot = [('sharecapT1Diff.png',Inches(6))],reco_text='The audit team shall investigate the capital activity files and the investor register files and determine the root cause of the variance.')

                    #share quantity testing
                    if (T1quantity_col==0).all():
                        workbook.Sheets('Share_Capital_T1').Select()
                        xl.ActiveWindow.Zoom = 70
                        xl.ActiveWindow.FreezePanes = False
                        workbook.Sheets('Share_Capital_T1').Cells(1,3).ColumnWidth = 24.33
                        workbook.Sheets('Share_Capital_T1').Cells(1,1).Select()
                        #in case it is in middle of woorkbook
                        time.sleep(1)
                        quantityzero = pyautogui.screenshot('quantityzero.png',region=(26,465, 1250, 485))
                        quantityzero = add_border(quantityzero,'quantityzero.png')

                        memo_final(memo_text ='Opening & closing share quantity and capital activity share quantity amount to zero.',screenshot = [('quantityzero.png',Inches(6))],reco_text='The audit team shall investigate the capital activity and investor register share quantity columns and determine the root cause of the variance.')

                    #capital activity detail columns
                    try:
                        CapAct_Detail = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_CapAct[0], sheet_name='Capital Activity', header=header_no)
                    except:
                        CapAct_DetailNotFound = True

                    if (abs(CapAct_Detail['Net Value Recalculation'])>1).any():
                        os.chdir(output_path+'/'+screenshot_path)
                        workbook.Sheets('Capital Activity').Select()
                        workbook.Sheets('Capital Activity').UsedRange.Find('Net Value Recalculation').Select()
                        xl.ActiveWindow.Zoom = 70
                        workbook.Sheets('Capital Activity').Cells(11,26).Select()
                        #in case it is in middle of woorkbook
                        time.sleep(1)
                        pyautogui.hotkey('ctrl', 'down')
                        pyautogui.press('down', presses=15)
                        time.sleep(1)
                        capactDetail = pyautogui.screenshot('capactDetail.png',region=(290,465, 820, 240))
                        capactDetail = add_border(capactDetail,'capactDetail.png')

                        memo_final(memo_text ='There is a difference in the Net Value Recalculation column of the Capital Activity tab in the Capital Activity working paper.',screenshot = [('capactDetail.png',Inches(5))],reco_text='The audit team shall investigate the capital activity input files and determine the root cause of the variance.')

                    workbook.Close(SaveChanges=False)
                    #xl.Application.Quit()

                ################CASH WORKPAPER TESTING#######################################################

                if len(WAMapps_CashRecon)!=0:
                    xl=client.gencache.EnsureDispatch("Excel.Application")
                    xl.Visible = True
                    xl.WindowState = client.constants.xlMaximized
                    xl.FormulaBarHeight = 2
                    workbook = xl.Workbooks.Open(output_path+'/'+wp_path+'/'+WAMapps_CashRecon[0])
                    workbook.Save()

                    try:
                        Cash_Summary = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_CashRecon[0], sheet_name='Summary', header=header_no, skipfooter=1)
                        Cash_temp1 = Cash_Summary.tail(2)
                        Cash_perTab = Cash_temp1.head(1)['Unnamed: 2'].iloc[0]
                        Cash_perLeads = Cash_temp1.tail(1)['Unnamed: 2'].iloc[0]
                        Cash_diff = (Cash_perTab-Cash_perLeads).round(2)
                        Cash_detail = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_CashRecon[0], sheet_name='Cash_Broker_Rec', header=header_no)
                    except:
                        CashNotFound = True

                    if abs(Cash_diff)>1:
                        cash_reco = []
                        reco_cash = []
                        cash_reclass_tab = []

                        cash_inputs = [(Cash_diff,dummy,'Cash','Cash','Cash_Reclass',Leads_Sheets,'Coversheet',cash_reco,cash_reclass_tab,reco_cash)]
                        recommendation(cash_inputs)

                        os.chdir(output_path+'/'+screenshot_path)
                        #navigate thru the workbook
                        workbook.Sheets('Summary').Select()
                        xl.ActiveWindow.Zoom = 70
                
                        workbook.Sheets('Summary').Cells(1,1).Select()
                        workbook.Sheets('Summary').Cells(1,3).ColumnWidth = 15
                        time.sleep(1)
                        CashNotZero = pyautogui.screenshot('CashNotZero.png',region=(52, 460,395 , 113))# Edited region=(50,450, 500, 150))
                        CashNotZero = add_border(CashNotZero,'CashNotZero.png')
                        

                        memo_final(memo_text ='Value per cash broker tab in cash working paper does not reconcile with C1 leads.',screenshot = [('CashNotZero.png',Inches(3))],reco_text=''.join(cash_reco))
                        cash_inputs = [(reco_cash,cash_reclass_tab,'cash')]
                        refer_memo(cash_inputs)


                    workbook.Close(SaveChanges=False)    
                    #xl.Application.Quit()

                ################DIVIDENDS TESTING#######################################################

                if len(WAMapps_Dividends)!=0:
                    xl=client.gencache.EnsureDispatch("Excel.Application")
                    xl.Visible = True
                    xl.WindowState = client.constants.xlMaximized
                    xl.FormulaBarHeight = 2
                    workbook = xl.Workbooks.Open(output_path+'/'+wp_path+'/'+WAMapps_Dividends[0])
                    workbook.Save()

                    try:
                        Div_Summary = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_Dividends[0], sheet_name='Summary', header=header_no, skipfooter=1)
                        with trial: Div_Inc_Tab = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_Dividends[0], sheet_name='Dividends_Income', header=header_no)
                        with trial: Div_Expense_Tab = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_Dividends[0], sheet_name='Dividends_Expense', header=header_no)
                        with trial: Div_Rec_Tab = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_Dividends[0], sheet_name='Dividends_Receivable', header=header_no)
                        with trial: Div_Pay_Tab = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_Dividends[0], sheet_name='Dividends_Payable', header=header_no)
                        Div_temp0 = Div_Summary[Div_Summary['Unnamed: 1']=='Difference'].fillna(0)
                        Div_temp1 = Div_Summary.tail(2)
                        Div_taxTab = Div_temp1.head(1).fillna(0)['Unnamed: 2'].iloc[0]
                        Div_taxLead = Div_temp1.tail(1).fillna(0)['Unnamed: 2'].iloc[0]
                        Div_Difference = Div_temp0['Unnamed: 2'].round(2)
                        Div_taxDiff = (Div_taxTab+Div_taxLead).round(2)
                        #for step2
                        DivIncome = Div_Summary[Div_Summary['Unnamed: 1']=='Value per Dividend_Income tab' ]['Unnamed: 2'].fillna(0).iloc[0]
                        DivExpense = Div_Summary[Div_Summary['Unnamed: 1']=='Value per Dividend_Expense tab' ]['Unnamed: 2'].fillna(0).iloc[0]
                        DivReceivable = Div_Summary[Div_Summary['Unnamed: 1']=='Value per Dividend_Receivable tab' ]['Unnamed: 2'].fillna(0).iloc[0]
                        DivPayable = Div_Summary[Div_Summary['Unnamed: 1']=='Value per Dividend_Payable tab' ]['Unnamed: 2'].fillna(0).iloc[0]
                        DivWHT = Div_Summary[Div_Summary['Unnamed: 1']=='Value per Dividend_Expense/Income tabs' ]['Unnamed: 2'].fillna(0).iloc[0]
                        #for detail tab variances
                        divIncBaseAmountRecalDiff = divIncBaseTaxRecalDiff = divExpBaseAmountRecalDiff = divExpBaseTaxRecalDiff = divRecBaseAmountRecalDiff = divRecBaseTaxRecalDiff = divPayBaseAmountRecalDiff = divPayBaseTaxRecalDiff = False
                        with trial: divIncBaseAmountRecalDiff = (abs(Div_Inc_Tab['Base Amount Recalculated Difference'])>1).any()
                        with trial: divIncBaseTaxRecalDiff = (abs(Div_Inc_Tab['Base Tax Recalculated Difference'])>1).any()
                        with trial: divExpBaseAmountRecalDiff = (abs(Div_Expense_Tab['Base Amount Recalculated Difference'])>1).any()
                        with trial: divExpBaseTaxRecalDiff = (abs(Div_Expense_Tab['Base Tax Recalculated Difference'])>1).any()
                        with trial: divRecBaseAmountRecalDiff = (abs(Div_Rec_Tab['Base Amount Recalculated Difference'])>1).any()
                        with trial: divRecBaseTaxRecalDiff = (abs(Div_Rec_Tab['Base Tax Recalculated Difference'])>1).any()
                        with trial: divPayBaseAmountRecalDiff = (abs(Div_Pay_Tab['Base Amount Recalculated Difference'])>1).any()
                        with trial: divPayBaseTaxRecalDiff = (abs(Div_Pay_Tab['Base Tax Recalculated Difference'])>1).any()
                        detail_tabs = []
                        column_tabs = []
                        inputs = []
                    except:
                        Div_NotFound = True

                    #Testing if any reconciliations in summdetail_tabsary have variance
                    if (abs(Div_Difference)>1).any() or abs(Div_taxDiff)>1:
                        os.chdir(output_path+'/'+screenshot_path)
                        #navigate thru the workbook
                        workbook.Sheets('Summary').Select()
                        workbook.Sheets('Summary').Cells(1,1).Select()
                        xl.ActiveWindow.FreezePanes = False
                        xl.ActiveWindow.Zoom = 60
                
                        pyautogui.press('down', presses=60)
                        time.sleep(1)
                        DivNotZero = pyautogui.screenshot('DivNotZero.png',region=(189,303,424,650))#Edited region=(90,305, 600, 665))
                        DivNotZero = add_border(DivNotZero,'DivNotZero.png')
                        

                        # to define the differences per mapping
                        div_names = []
                        divreco = []
                        reco_divinc = []
                        reco_divexp = []
                        reco_divrec = []
                        reco_divpay = []
                        reco_divwht = []
                        divincome_reclass_tab = []
                        divexp_reclass_tab = []
                        divrec_reclass_tab = []
                        divpay_reclass_tab = [] 
                        divtax_reclass_tab = []
                        div_inputs = []
                        Div_Sheets = pd.ExcelFile(output_path+'/'+wp_path+'/'+WAMapps_Dividends[0])
                        with trial: div_inputs.append((Div_Difference.iloc[0],div_names,'Dividend Income','DivIncome','DivIncome_Reclass',Div_Sheets,'Dividends_Income',divreco,divincome_reclass_tab,reco_divinc))
                        with trial: div_inputs.append((Div_Difference.iloc[1],div_names,'Dividend Expense','DivExpense','DivExpense_Reclass',Div_Sheets,'Dividends_Expense',divreco,divexp_reclass_tab,reco_divexp))
                        with trial: div_inputs.append((Div_Difference.iloc[2],div_names,'Dividend Receivable','DivRec','DivRec_Reclass',Div_Sheets,'Dividends_Receivable',divreco,divrec_reclass_tab,reco_divrec))
                        with trial: div_inputs.append((Div_Difference.iloc[3],div_names,'Dividend Payable','DivPay','DivPay_Reclass',Div_Sheets,'Dividends_Payable',divreco,divpay_reclass_tab,reco_divpay))
                        with trial: div_inputs.append((Div_taxDiff,div_names,'Dividend Withholding Tax','DivWHT','DivWHT_Reclass',Div_Sheets,'Summary',divreco,divtax_reclass_tab,reco_divwht))

                        recommendation(div_inputs)

                        #string conversion
                        div_names_str = ', '.join(div_names)
                        div_reco_str = ' \n\n'.join(divreco)


                        ##for proper preposition and grammar
                        if len(div_names)>1:
                            preposition = 'do not reconcile with their Leads counterparts'
                        else:
                            preposition = 'does not reconcile with its Leads counterpart'


                        memo_final(memo_text ='%s per dividends working paper %s.'%(div_names_str,preposition),screenshot = [('DivNotZero.png',Inches(2.5))],reco_text=div_reco_str)
                        dividends_inputs = [(reco_divinc,divincome_reclass_tab,'dividend income'),(reco_divexp,divexp_reclass_tab,'dividend expense'),(reco_divrec,divrec_reclass_tab,'dividend receivable'),(reco_divpay,divpay_reclass_tab,'dividend payable'),(reco_divwht,divtax_reclass_tab,'dividend WHT')]
                        refer_memo(dividends_inputs)

                    ##detail tab testing

                    if divIncBaseAmountRecalDiff or divIncBaseTaxRecalDiff or divExpBaseAmountRecalDiff or divExpBaseTaxRecalDiff or divRecBaseAmountRecalDiff or divRecBaseTaxRecalDiff or divPayBaseAmountRecalDiff or divPayBaseTaxRecalDiff:
                        with trial: inputs.append((divIncBaseAmountRecalDiff,divIncBaseTaxRecalDiff,'Dividend Income','Dividends_Income'))
                        with trial: inputs.append((divExpBaseAmountRecalDiff,divExpBaseTaxRecalDiff,'Dividend Expense','Dividends_Expense'))
                        with trial: inputs.append((divRecBaseAmountRecalDiff,divRecBaseTaxRecalDiff,'Dividend Receivable','Dividends_Receivable'))
                        with trial: inputs.append((divPayBaseAmountRecalDiff,divPayBaseTaxRecalDiff,'Dividend Payable','Dividends_Payable'))
                        for (value1,value2,append_name,tab) in inputs:    
                            if value1 or value2:
                                detail_tabs.append(append_name)
                                if value1:
                                    column_tabs.append('Base Amount Recalculated Difference')
                                if value2:
                                    column_tabs.append('Base Tax Recalculated Difference')
                                os.chdir(output_path+'/'+screenshot_path)
                                #navigate thru the workbook
                                workbook.Sheets(tab).Select()
                                workbook.Sheets(tab).UsedRange.Find('Base Amount Recalculated Difference').Select()
                                xl.ActiveWindow.Zoom = 70
                                pyautogui.hotkey('ctrl','down')
                                pyautogui.press('down',presses=17)
                                time.sleep(1)
                                DivDetailNotZero = pyautogui.screenshot('DivDetailNotZero.png',region=(813,465,598,208))
                                DivDetailNotZero = add_border(DivDetailNotZero,'DivDetailNotZero.png')
                        detail_tabs_str = ', '.join(detail_tabs)
                        column_tabsSum = list(set(column_tabs))
                        column_tabsSum_str = ', '.join(column_tabsSum)

                        memo_final(memo_text ='There is a difference in the %s columns of the %s tabs in the dividends working paper.'%(column_tabsSum_str, detail_tabs_str),screenshot = [('DivDetailNotZero.png',Inches(4.5))],reco_text = 'The audit team shall investigate the dividend input files and investigate the root cause of the issue.')

                    workbook.Close(SaveChanges=False)    
                    #xl.Application.Quit()

                ################INTEREST TESTING#######################################################

                if len(WAMapps_Interest)!=0:
                    xl=client.gencache.EnsureDispatch("Excel.Application")
                    xl.Visible = True
                    xl.WindowState = client.constants.xlMaximized
                    xl.FormulaBarHeight = 2
                    workbook = xl.Workbooks.Open(output_path+'/'+wp_path+'/'+WAMapps_Interest[0])
                    workbook.Save()

                    try:
                        Int_Summary = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_Interest[0], sheet_name='Summary', header=25, skipfooter=1)
                        with trial: Int_Inc_Tab = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_Interest[0], sheet_name='Interest_Income', header=header_no)
                        with trial: Int_Expense_Tab = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_Interest[0], sheet_name='Interest_Expense', header=header_no)
                        with trial: Int_Rec_Tab = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_Interest[0], sheet_name='Interest_Receivable', header=header_no)
                        with trial: Int_Pay_Tab = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_Interest[0], sheet_name='Interest_Payable', header=header_no)
                        Int_temp0 = Int_Summary[Int_Summary['Unnamed: 1']=='Difference'].fillna(0)
                        Int_temp1 = Int_Summary.tail(2)
                        Int_taxTab = Int_temp1.head(1).fillna(0)['Unnamed: 2'].iloc[0]
                        Int_taxLead = Int_temp1.tail(1).fillna(0)['Unnamed: 2'].iloc[0]
                        Int_Difference = Int_temp0['Unnamed: 2'].round(2)
                        Int_taxDiff = (Int_taxTab-Int_taxTab).round(2)
                        #for step2
                        IntIncome = Int_Summary[Int_Summary['Unnamed: 1']=='Value per Interest_Income tab' ]['Unnamed: 2'].fillna(0).iloc[0]
                        IntExpense = Int_Summary[Int_Summary['Unnamed: 1']=='Value per Interest_Expense tab' ]['Unnamed: 2'].fillna(0).iloc[0]
                        IntReceivable = Int_Summary[Int_Summary['Unnamed: 1']=='Value per Interest_Receivable tab' ]['Unnamed: 2'].fillna(0).iloc[0]
                        IntPayable = Int_Summary[Int_Summary['Unnamed: 1']=='Value per Interest_Payable tab' ]['Unnamed: 2'].fillna(0).iloc[0]
                        IntWHT = Int_Summary[Int_Summary['Unnamed: 1']=='Value per Interest_Expense/Income tabs' ]['Unnamed: 2'].fillna(0).iloc[0]        
                    except:
                        Int_NotFound = True

                    #Testing if any reconciliations in summary have variance
                    if (abs(Int_Difference)>1).any() or abs(Int_taxDiff)>1:
                        os.chdir(output_path+'/'+screenshot_path)
                        #navigate thru the workbook
                        workbook.Sheets('Summary').Select()
                        workbook.Sheets('Summary').Cells(1,1).Select()
                        xl.ActiveWindow.FreezePanes = False
                        xl.ActiveWindow.Zoom = 60
                
                        pyautogui.press('down', presses=59)
                        pyautogui.press('right', presses=5) 
                        time.sleep(1)
                        IntNotZero = pyautogui.screenshot('IntNotZero.png',region=(189,302,415,653))#Edited region=(90,310, 600, 680))
                        IntNotZero = add_border(IntNotZero,'IntNotZero.png')
                        

                        #to define the differences per mapping
                        int_names = []
                        intreco = []
                        reco_intinc = []
                        reco_intexp = []
                        reco_intrec = []
                        reco_intpay = []
                        reco_intwht = []
                        intinc_reclass_tab = [] 
                        intexp_reclass_tab = [] 
                        intrec_reclass_tab = [] 
                        intpay_reclass_tab = [] 
                        inttax_reclass_tab = []
                        int_inputs = []
                        Int_Sheets = pd.ExcelFile(output_path+'/'+wp_path+'/'+WAMapps_Interest[0])
                        with trial: int_inputs.append((Int_Difference.iloc[0],int_names,'Interest Income','IntIncome','IntIncome_Reclass',Int_Sheets,'Interest_Income',intreco,intinc_reclass_tab,reco_intinc))
                        with trial: int_inputs.append((Int_Difference.iloc[1],int_names,'Interest Expense','IntExpense','IntExpense_Reclass',Int_Sheets,'Interest_Expense',intreco,intexp_reclass_tab,reco_intexp))
                        with trial: int_inputs.append((Int_Difference.iloc[2],int_names,'Interest Receivable','IntRec','IntRec_Reclass',Int_Sheets,'Interest_Receivable',intreco,intrec_reclass_tab,reco_intrec))
                        with trial: int_inputs.append((Int_Difference.iloc[3],int_names,'Interest Payable','IntPay','IntPay_Reclass',Int_Sheets,'Interest_Payable',intreco,intpay_reclass_tab,reco_intpay))
                        with trial: int_inputs.append((Int_taxDiff,int_names,'Interest Withholding Tax','IntWHT','IntWHT_Reclass',Int_Sheets,'Summary',intreco,inttax_reclass_tab,reco_intwht))

                        recommendation(int_inputs)    

                        int_names_str = ', '.join(int_names)
                        int_reco_str = ' \n\n'.join(intreco)

                        ##for proper preposition and grammar
                        if len(int_names)>1:
                            preposition = 'do not reconcile with their Leads counterparts'
                        else:
                            preposition = 'does not reconcile with its Leads counterpart'

                        memo_final(memo_text ='%s per interest working paper %s.'%(int_names_str,preposition),screenshot = [('IntNotZero.png',Inches(2.5))],reco_text=int_reco_str)
                        interest_inputs = [(reco_intinc,intinc_reclass_tab,'interest income'),(reco_intexp,intexp_reclass_tab,'interest expense'),(reco_intrec,intrec_reclass_tab,'interest receivable'),(reco_intpay,intpay_reclass_tab,'interest payable'),(reco_intwht,inttax_reclass_tab,'interest WHT')]
                        refer_memo(interest_inputs)

                    workbook.Close(SaveChanges=False)    
                    #xl.Application.Quit()


                #####FX Reasonability Testing ##############################

                if len(WAMapps_FXReason)!=0:
                    xl=client.gencache.EnsureDispatch("Excel.Application")
                    xl.Visible = True
                    xl.WindowState = client.constants.xlMaximized
                    xl.FormulaBarHeight = 2
                    workbook = xl.Workbooks.Open(output_path+'/'+wp_path+'/'+WAMapps_FXReason[0])
                    workbook.Save()

                    try:
                        FX_Summary = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_FXReason[0], sheet_name='FX_Reasonability', header=header_no+4)
                        FX_items_df = FX_Summary[FX_Summary['Data Source']!=''].fillna('None')
                        FX_items = FX_items_df['FX Rate'].tolist()
                        FX_dtypes = []
                        for i in range(0, len(FX_items)):
                            FX_dtypes.append(str(type(FX_items[i])))
                        FX_items_df['dtypes']=FX_dtypes
                        dtypes_cond = (~FX_items_df['dtypes'].str.contains('int') & ~FX_items_df['dtypes'].str.contains('float')).to_list()
                        FX_items_df['condition']=dtypes_cond

                        #for list paragraph purposes    
                        blank_FX = FX_items_df[(FX_items_df['condition']==True) & (FX_items_df['FX Rate']=='None')]['Data Source'].to_list()
                        invalid_FX = FX_items_df[(FX_items_df['condition']==True) & (FX_items_df['FX Rate']!='None')]['Data Source'].to_list()
                        blank_FX = list(set(blank_FX))
                        invalid_FX = list(set(invalid_FX))

                        #for grammar purposes
                        blank_FX_final = []
                        for word in blank_FX:
                            word2 = word.replace('Balance Sheet - ','')
                            blank_FX_final.append(word2)

                        invalid_FX_final = []
                        for word in invalid_FX:
                            word2 = word.replace('Balance Sheet - ','')
                            invalid_FX_final.append(word2)

                        blank_FX_str = ', '.join(blank_FX_final).title()
                        invalid_FX_str = ', '.join(invalid_FX_final).title()

                        if len(blank_FX)==1:            
                            blank_FX_text = '%s data source has a missing FX rate as per FX reasonability working paper.'%(blank_FX_str)
                        elif len(blank_FX)>1: 
                            blank_FX_text = '%s data sources have missing FX rates as per FX reasonability working paper.'%(blank_FX_str)
                        else:
                            pass

                        if len(invalid_FX)==1:
                            invalid_FX_text = '%s data source has invalid FX rate as per FX reasonability working paper.'%(invalid_FX_str)
                        elif len(invalid_FX)>1:
                            invalid_FX_text = '%s data sources have invalid FX rates as per FX reasonability working paper.'%(invalid_FX_str)
                        else:
                            pass
                    except:
                        FX_NotFound = True

                    ###to automate the text in case of invalid/blank FX rate
                    if len(invalid_FX)!=0 and len(blank_FX)!=0:
                        FX_text = blank_FX_text+' '+invalid_FX_text
                    elif len(invalid_FX)!=0 and len(blank_FX)==0:
                        FX_text = invalid_FX_text
                    elif len(invalid_FX)==0 and len(blank_FX)!=0:
                        FX_text = blank_FX_text
                    else:
                        pass

                    if (FX_items_df['condition']==True).any():
                        os.chdir(output_path+'/'+screenshot_path)
                        #navigate thru the workbook
                        workbook.Sheets('FX_Reasonability').Select()
                        xl.ActiveWindow.Zoom = 70
                        workbook.Sheets('FX_Reasonability').Cells(15,2).Select()
                        workbook.Sheets('FX_Reasonability').Cells(1,3).ColumnWidth = 24
                
                        #Sorting alphabetically
                        pyautogui.hotkey('alt','a','s','a')
                        #For row height (region length)
                        fx_height = len(FX_items_df)+1
                        #Limiting screenshot length if too much rates
                        if fx_height >= 21:
                            fx_height = 21 * 18
                        elif fx_height <= 7:
                            fx_height = fx_height * 22
                        else:
                            fx_height = fx_height * 18
                        time.sleep(1)
                        FXblank = pyautogui.screenshot('FXblank.png',region=(84,527,919,fx_height))#Edited region=(60,305, 920, 665))
                        FXblank = add_border(FXblank,'FXblank.png')
                        
                        memo_final(memo_text = FX_text,screenshot = [('FXblank.png',Inches(6))],reco_text='The audit team shall manually input the missing FX rates as this is a data limitation issue.')

                    workbook.Close(SaveChanges=False)    
                    #xl.Application.Quit()

                #####Investment Rollforward Testing ##############################

                if len(WAMapps_InvRF)!=0:
                    xl=client.gencache.EnsureDispatch("Excel.Application")
                    xl.Visible = True
                    xl.WindowState = client.constants.xlMaximized
                    xl.FormulaBarHeight = 2
                    workbook = xl.Workbooks.Open(output_path+'/'+wp_path+'/'+WAMapps_InvRF[0])
                    workbook.Save()

                    try:
                        IRF_Summary = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_InvRF[0], sheet_name='Summary', header=header_no, skipfooter=1)
                        IRF_temp0 = IRF_Summary[IRF_Summary['CY Portfolio Rec']=='Difference'].fillna(0)
                        IRF_temp1 = IRF_Summary.tail(2)
                        IRF_PyTab = IRF_temp1.head(1).fillna(0)['Unnamed: 2'].iloc[0]
                        IRF_PyLead = IRF_temp1.tail(1).fillna(0)['Unnamed: 2'].iloc[0]
                        IRF_Difference = IRF_temp0['Unnamed: 2'].round(2)
                        IRF_PyDiff = (IRF_PyTab-IRF_PyLead).round(2)
                        #FOR STEP 2
                        ClosingFV = IRF_Summary[IRF_Summary['CY Portfolio Rec']=='Closing Fair Value per Roll_Forward tab']['Unnamed: 2'].fillna(0).iloc[0]
                        OpeningFV = IRF_Summary[IRF_Summary['CY Portfolio Rec']=='Opening Fair Value per Roll_Forward tab']['Unnamed: 2'].fillna(0).iloc[0]
                        RGL_IRF = IRF_Summary[IRF_Summary['CY Portfolio Rec']=='Realized Gain (Loss) per Roll_Forward tab']['Unnamed: 2'].fillna(0).iloc[0]
                        ClosingCost = IRF_Summary[IRF_Summary['CY Portfolio Rec']=='Closing Cost per Roll_Forward tab']['Unnamed: 2'].fillna(0).iloc[0]
                        OpeningCost = IRF_Summary[IRF_Summary['CY Portfolio Rec']=='Opening Cost per Roll_Forward tab']['Unnamed: 2'].fillna(0).iloc[0]
                        #quantity and cost RF testing
                        IRF_Rollforward = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_InvRF[0], sheet_name='Roll_Forward', header=header_no+1)
                        #quantity testing:
                        IRF_quantity = IRF_Rollforward['Difference']
                        #cost testing:
                        IRF_cost = IRF_Rollforward['Rec']
                    except:
                        IRF_NotFound = True

                    #Testing if any reconciliations in summary have variance
                    if (abs(IRF_Difference)>1).any() or abs(IRF_PyDiff)>1:
                        os.chdir(output_path+'/'+screenshot_path)
                        #navigate thru the workbook
                        workbook.Sheets('Summary').Select()
                        workbook.Sheets('Summary').Cells(1,1).Select()
                        xl.ActiveWindow.FreezePanes = False
                        xl.ActiveWindow.Zoom = 60
                        #in case it is in middle of woorkbook
                
                        pyautogui.press('down', presses=56)
                        workbook.Sheets('Summary').Rows(18).EntireRow.Delete()
                        workbook.Sheets('Summary').Rows(26).EntireRow.Delete()
                        workbook.Sheets('Summary').Rows(34).EntireRow.Delete()
                        workbook.Sheets('Summary').Rows(42).EntireRow.Delete()
                        workbook.Sheets('Summary').Rows(50).EntireRow.Delete()
                        workbook.Sheets('Summary').Cells(1,3).ColumnWidth = 24
                        time.sleep(1)
                        IRFNotZero = pyautogui.screenshot('IRFNotZero.png',region=(75,305, 520, 665))
                        IRFNotZero = add_border(IRFNotZero,'IRFNotZero.png')
                        

                        # to define the differences per mapping
                        irf_names = []
                        irf_reco = []
                        reco_closingFV = []
                        reco_openingFV = []
                        reco_RGL = []
                        reco_UGL = []
                        reco_closingCost = []
                        reco_openingCost = []
                        fvcy_reclass_tab = []
                        fvpy_reclass_tab = [] 
                        rgl_reclass_tab = [] 
                        ugl_reclass_tab = [] 
                        irfcost_reclass_tab = []
                        irfcostpy_reclass_tab = []

                        irf_inputs = [(IRF_Difference.iloc[0],irf_names,'Closing fair value','IRF_FV_CY','IRF_FV_CY_Reclass',Leads_Sheets,'Coversheet',irf_reco,fvcy_reclass_tab,reco_closingFV)
                                    ,(IRF_Difference.iloc[1],irf_names,'Opening fair value','IRF_FV_PY','IRF_FV_PY_Reclass',Leads_Sheets,'Coversheet',irf_reco,fvpy_reclass_tab,reco_openingFV)
                                    ,(IRF_Difference.iloc[2],irf_names,'RGL','RGL_IRF_CY','RGL_IRF_CY_Reclass',Leads_Sheets,'Coversheet',irf_reco,rgl_reclass_tab,reco_RGL)
                                    ,(IRF_Difference.iloc[3],irf_names,'UGL','UGL_IRF_CY','UGL_IRF_CY_Reclass',Leads_Sheets,'Coversheet',irf_reco,ugl_reclass_tab,reco_UGL)
                                    ,(IRF_Difference.iloc[4],irf_names,'Closing cost','IRF_Cost_CY','IRF_Cost_CY_Reclass',Leads_Sheets,'Coversheet',irf_reco,irfcost_reclass_tab,reco_closingCost)
                                    ,(IRF_PyDiff,irf_names,'Opening cost','IRF_Cost_PY','IRF_Cost_PY_Reclass',Leads_Sheets,'Coversheet',irf_reco,irfcostpy_reclass_tab,reco_openingCost)]

                        recommendation(irf_inputs)

                        irf_names_str = ', '.join(irf_names)
                        irf_reco_str = ' \n\n'.join(irf_reco)

                        if len(irf_names)>1:
                            preposition = 'do not reconcile with their Leads counterparts'
                        else:
                            preposition = 'does not reconcile with its Leads counterpart'


                        memo_final(memo_text ='%s per IRF working paper %s.'%(irf_names_str,preposition),screenshot = [('IRFNotZero.png',Inches(2.5))],reco_text=irf_reco_str)
                        irf_inputs = [(reco_closingFV,fvcy_reclass_tab,'closing portfolio FV')
                                    ,(reco_openingFV,fvpy_reclass_tab,'opening portfolio FV')
                                    ,(reco_RGL,rgl_reclass_tab,'RGL')
                                    ,(reco_UGL,ugl_reclass_tab,'UGL')
                                    ,(reco_closingCost,irfcost_reclass_tab,'closing portfolio cost')
                                    ,(reco_openingCost,irfcostpy_reclass_tab,'opening portfolio cost')]

                        refer_memo(irf_inputs)


                #quantity and cost testing    
                    if (abs(IRF_quantity)>1).any() or (abs(IRF_cost)>1).any():
                        os.chdir(output_path+'/'+screenshot_path)

                        Pivot = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_InvRF[0], sheet_name='BCR Pivot', header=9)
                        Pivot_height = len(Pivot)
                        workbook.Sheets('BCR Pivot').Select()
                        workbook.Sheets('BCR Pivot').Cells(1,1).Select()
                        workbook.Sheets('BCR Pivot').Cells(1,4).ColumnWidth = 25
                        workbook.Sheets('BCR Pivot').Cells(1,5).ColumnWidth = 17
                        workbook.Sheets('BCR Pivot').Cells(1,6).ColumnWidth = 17
                        workbook.Sheets('BCR Pivot').Cells(1,7).ColumnWidth = 17
                        workbook.Sheets('BCR Pivot').Cells(1,9).ColumnWidth = 17
                        workbook.Sheets('BCR Pivot').Cells(1,10).ColumnWidth = 17
                        workbook.Sheets('BCR Pivot').Cells(1,11).ColumnWidth = 17
                        workbook.Sheets('BCR Pivot').Cells(1,1).Select()
                        nrows = workbook.Sheets('BCR Pivot').UsedRange.Rows.Count
                        Pivot_height = nrows * 12
                        if Pivot_height > 500:
                            Pivot_height = 500
                        else:
                            pass    
                        time.sleep(1)
            
                        PivotNotZero = pyautogui.screenshot('PivotNotZero.png',region=(57,444,1194,Pivot_height))#Edited region=(40,305, 1700, 600))
                        PivotNotZero = add_border(PivotNotZero,'PivotNotZero.png')

                        workbook.Sheets('Roll_Forward').Select()

                        #For screenshot height
                        IRF_quantity_len = IRF_quantity[abs(IRF_Rollforward['Difference'])>1].count()
                        if IRF_quantity_len <= 6:
                            quantity_height = 130
                        elif IRF_quantity_len <= 11:
                            quantity_height = 200
                        else:
                            quantity_height = 250

                        #For screenshot height
                        IRF_cost_len = IRF_cost[abs(IRF_Rollforward['Rec'])>1].count()
                        if IRF_cost_len <= 6:
                            cost_height = 130
                        elif IRF_cost_len <= 11:
                            cost_height = 200
                        else:
                            cost_height = 250

                        #Quantity Screenshot
                        workbook.Sheets('Roll_Forward').Cells(12,2).Select()
                        xl.ActiveWindow.Zoom = 60
                        pyautogui.press('right',presses=21)
                        pyautogui.press('left',presses=7)
                        # pyautogui.hotkey('alt', 'down')
                        # pyautogui.hotkey('f', 'g','1','tab','right','tab')
                        # pyautogui.press('down',presses=6)
                        # pyautogui.hotkey('enter','tab','-','1','enter')
                        # time.sleep(2)
                        pyautogui.hotkey('ctrl','down')
                        pyautogui.press('down',presses=21)
                        time.sleep(1)
                        QuantityNotZero = pyautogui.screenshot('QuantityNotZero.png',region=(15,430,835,quantity_height))#Edited egion=(10,305, 1700, 600))
                        QuantityNotZero = add_border(QuantityNotZero,'QuantityNotZero.png')

                        #Cost Screenshot
                        workbook.Sheets('Roll_Forward').Cells(12,2).Select()
                        xl.ActiveWindow.Zoom = 60
                        pyautogui.hotkey('alt', 'a', 't')
                        pyautogui.hotkey('alt', 'a', 't')
                        pyautogui.press('right',presses=21)
                        # pyautogui.hotkey('alt', 'down')
                        # pyautogui.hotkey('f', 'g','1','tab','right','tab')
                        # pyautogui.press('down',presses=6)
                        # pyautogui.hotkey('enter','tab','-','1','enter')
                        # time.sleep(2)
                        pyautogui.hotkey('ctrl','down')
                        pyautogui.press('down',presses=21)
                        time.sleep(1)
                        CostNotZero = pyautogui.screenshot('CostNotZero.png',region=(825,427,972,cost_height))#Edited region=(5,305, 1700, 600))
                        CostNotZero = add_border(CostNotZero,'CostNotZero.png')
                        pyautogui.hotkey('alt', 'a', 't')

                        if (abs(IRF_quantity)>1).any() and (abs(IRF_cost)>1).any():
                            sec3 = 'Quantity and cost rollforward in the IRF working paper do not reconcile.'
                            irf_ss = [('QuantityNotZero.png',Inches(5)),('CostNotZero.png',Inches(5)),('PivotNotZero.png',Inches(5))]
                        elif (abs(IRF_quantity)>1).any():
                            sec3 = 'Quantity rollforward in the IRF working paper does not reconcile.'
                            irf_ss = [('QuantityNotZero.png',Inches(5))]
                        elif (abs(IRF_cost)>1).any():
                            sec3 = 'Cost rollforward in the IRF working paper does not reconcile.'
                            irf_ss = [('CostNotZero.png',Inches(5)),('PivotNotZero.png',Inches(5))]
                        else:
                            pass

                        memo_final(memo_text = sec3,screenshot = irf_ss,reco_text='The audit team shall investigate the PY & CY portfolio, purchases & sales and RGL input files and determine the reason for the variance.')        

                    else:
                        pass


                    ###missing mapping testing
                    if (IRF_Rollforward[(IRF_Rollforward['Investment Class'].str.lower()=='missing') & (IRF_Rollforward['Investment Type'].isnull())]['Fund ID'].notnull()).any():
                        os.chdir(output_path+'/'+screenshot_path)
                        #navigate thru the workbook
                
                        workbook.Sheets('Roll_Forward').Select()
                        workbook.Sheets('Roll_Forward').Cells(1,1).Select()
                        workbook.Sheets('Roll_Forward').Cells(20,20).Select()
                        try:
                            workbook.Sheets('Roll_Forward').Range('I:I').Find('Missing').Select()
                        except:
                            workbook.Sheets('Roll_Forward').Range('I:I').Find('MISSING').Select()
                        xl.ActiveWindow.Zoom = 60
                        workbook.Sheets('Roll_Forward').Cells(1,4).ColumnWidth = 24
                        workbook.Sheets('Roll_Forward').Cells(1,5).ColumnWidth = 24
                        workbook.Sheets('Roll_Forward').Cells(1,6).ColumnWidth = 24
                        workbook.Sheets('Roll_Forward').Cells(1,7).ColumnWidth = 53
                        workbook.Sheets('Roll_Forward').Cells(1,8).ColumnWidth = 39
                        workbook.Sheets('Roll_Forward').Cells(1,9).ColumnWidth = 24
                        workbook.Sheets('Roll_Forward').Cells(1,10).ColumnWidth = 24
                        time.sleep(1)
                        IRFMissingMapping = pyautogui.screenshot('IRFMissingMapping.png',region=(27,443,1175,530))#Edited region=(10,305, 1700, 600))
                        IRFMissingMapping = add_border(IRFMissingMapping,'IRFMissingMapping.png')
                        

                        #recommendation section
                        if 'alpha' in doc_fr['TPAName'][row_num]:
                            reco15a = 'This is due to data limitation. The said missing mappings are those investments that are present in the “Purchases and Sales” file and missing in “Cost Roll Forward” file. They use “InvestID” and “Investment” as the common column, respectively, for the former to obtain the “Group2” column from the latter which is used as the common field for acquiring the instrument mapping.'
                        else:
                            reco15a = 'This is due to data limitation. The value in the Investment Type column equivalent in the raw file is blank as such instrument was not mapped.'

                        memo_final(memo_text ='There are instruments with missing mapping in the IRF workpaper.',screenshot = [('IRFMissingMapping.png',Inches(6))],reco_text=reco15a)


                    workbook.Close(SaveChanges=False)    
                    #xl.Application.Quit()

                #####Portfolio Valuations Testing ##############################

                if len(WAMapps_PortVal)!=0:
                    xl=client.gencache.EnsureDispatch("Excel.Application")
                    xl.Visible = True
                    xl.WindowState = client.constants.xlMaximized
                    xl.FormulaBarHeight = 2
                    workbook = xl.Workbooks.Open(output_path+'/'+wp_path+'/'+WAMapps_PortVal[0])
                    workbook.Save()

                    ##Tie to D1 testing
                    try:
                        PortVal_TietoD1 = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_PortVal[0], sheet_name='Tie_To_D1', header=header_no+1, skipfooter=1)
                    except:
                        PortVal_NotFound = True

                    if (abs(PortVal_TietoD1['Difference'].fillna(0))>1).any():
                        os.chdir(output_path+'/'+screenshot_path)
                        #navigate thru the workbook
                        workbook.Sheets('Tie_To_D1').Select()
                        xl.ActiveWindow.Zoom = 70
                        xl.ActiveWindow.FreezePanes = False
                        workbook.Sheets('Tie_To_D1').Cells(1,1).Select()
                
                        #Height
                        Port_height = (PortVal_TietoD1['Difference'].fillna(0).count()+3)
                        
                        if Port_height <= 15:
                            Port_height = Port_height * 22
                        elif Port_height <= 35:
                            Port_height = Port_height * 18
                        else:
                            Port_height = 630
                            # Port_height = Port_height * 18
                    
                        pyautogui.press('down', presses=47)
                        workbook.Sheets('Tie_To_D1').Cells(1,2).ColumnWidth = 24
                        workbook.Sheets('Tie_To_D1').Cells(1,3).ColumnWidth = 29
                        workbook.Sheets('Tie_To_D1').Cells(1,5).ColumnWidth = 18
                        workbook.Sheets('Tie_To_D1').Cells(1,7).ColumnWidth = 18
                        time.sleep(1)
                        PortNotZero = pyautogui.screenshot('PortNotZero.png',region=(50,300,1100, Port_height))
                        PortNotZero = add_border(PortNotZero,'PortNotZero.png')
                        if (abs(PortVal_TietoD1['Difference'].fillna(0).sum())<1):
                            reco_tietod1 = 'This is a mapping issue considering that the totals have net to nil. As such, the audit team may opt to manually reclassify the mapping of the instruments or they may update the account/instrument mapping file and ask WAMapps team for a refreshed output.'
                        else:
                            reco_tietod1 = 'The audit team shall investigate the portfolio input file and determine the root cause of the variance.'
                        memo_final(memo_text ='Fair value per portfolio working paper does not reconcile with D1 Leads.',screenshot = [('PortNotZero.png',Inches(6))],reco_text=reco_tietod1)

                    ##Pricing Summary testing
                    try:
                        PortVal_Pricing = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_PortVal[0], sheet_name='Pricing_Summary', header=22)
                        UnpricedItems = PortVal_Pricing[PortVal_Pricing['Unnamed: 1']=='Number of unpriced items']['Unnamed: 2'].iloc[0]
                        pricedItems = PortVal_Pricing[PortVal_Pricing['Unnamed: 1']=='Number of beta priced items']['Unnamed: 2'].iloc[0]
                        pricedDifference = abs(PortVal_Pricing[PortVal_Pricing['Unnamed: 1']=='Difference (%)']['Unnamed: 2'].iloc[0])
                    except:
                        PortVal_PricingNotFound = True

                    #Screenshot
                    os.chdir(output_path+'/'+screenshot_path)
                    #navigate thru the workbook
                    workbook.Sheets('Pricing_Summary').Select()
                    xl.ActiveWindow.Zoom = 70
                    workbook.Sheets('Pricing_Summary').Cells(1,1).Select()
                    xl.ActiveWindow.FreezePanes = False
                    workbook.Sheets('Pricing_Summary').Cells(1,3).ColumnWidth = 20
                    workbook.Sheets('Pricing_Summary').Cells(1,4).ColumnWidth = 10
                    time.sleep(1)
                    PricingNotZero = pyautogui.screenshot('PricingNotZero.png',region=(150,440, 500, 500)) #Edited region=(100,445, 700, 500))
                    PricingNotZero = add_border(PricingNotZero,'PricingNotZero.png')
                    #move to beta alpha worksheet
                    try:
                        workbook.Sheets('alpha Price Support').Select()
                        workbook.Sheets('alpha Price Support').Cells(1,1).Select()
                        time.sleep(1)
                    except:
                        pass
                    alphass = pyautogui.screenshot('alpha.png',region=(50,310, 1550, 350))
                    alphass = add_border(alphass,'alpha.png')

                    #Pricing Issue Memo Logic
                    if abs(UnpricedItems)>0 or abs(pricedDifference)>=.10:
                        port_memo_wdiff = 'The audit team shall investigate the ticker columns of the portfolio input file as they might be invalid in format. The case may also be that these are private instruments as such the audit team shall manually price them. \n \nThe reason for the difference is might be due to the following:\n•The instruments may need to be scaled, as such audit team may opt to tag as YES the Scale by 100 column of the Portfolio tab in the Portfolio working paper\n•Incorrect or missing FX rates\n•The beta Prices do not agree with the Client price'
                        port_memo_wodiff = 'The audit team shall investigate the ticker columns of the portfolio input file as they might be invalid in format. The case may also be that these are private instruments as such the audit team shall manually price them.'
                        #Valuation is N
                        if doc_fr['External Valuations'][row_num]=='N':
                            sec3 = 'No external pricing was pulled from the alpha portal for the portfolio valuations workpaper due to "External Valuations" in the WAMapps request tagged as "N".'
                            port_ss = [('PricingNotZero.png',Inches(2.5))]
                            port_memo = 'Audit team should confirm if they do not really need the pricing. Should they need the valuation, they will need to request a rerun of the Portfolio Valuations to the WAMapps delivery team.'
                        #No Priced Items
                        elif abs(pricedItems)==0:
                            sec3 = 'No external pricing was pulled from the alpha portal for the portfolio valuations workpaper'
                            port_ss = [('PricingNotZero.png',Inches(2.5)),('alpha.png',Inches(4.5))]
                            port_memo = port_memo_wodiff
                        #Some Priced Items
                        elif abs(UnpricedItems)>0 and abs(pricedDifference)>.10:
                            sec3 = 'Some instruments were not externally priced in the alpha portal and there was a difference noted between the fair value of price items per client and per beta in the portfolio valuations workpaper.'
                            port_ss = [('PricingNotZero.png',Inches(2.5))]
                            port_memo = port_memo_wdiff
                        elif abs(UnpricedItems)>0:
                            sec3 = 'Some instruments were not externally priced in the alpha portal for the portfolio valuations workpaper.'
                            port_ss = [('PricingNotZero.png',Inches(2.5))]
                            port_memo = port_memo_wodiff
                        #All Items Priced but with difference
                        elif abs(UnpricedItems)==0 and abs(pricedDifference)>.10:
                            sec3 = 'All instruments were externally priced in the alpha portal for the portfolio valuations workpaper however there was a difference noted between the fair value of price items per client and per beta.'
                            port_ss = [('PricingNotZero.png',Inches(2.5))]
                            port_memo = 'The reason for the difference is might be due to the following:\n•The instruments may need to be scaled, as such audit team may opt to tag as YES the Scale by 100 column of the Portfolio tab in the Portfolio working paper\n•Incorrect or missing FX rates\n•The beta Prices do not agree with the Client price'

                        memo_final(memo_text = sec3,screenshot = port_ss,reco_text=port_memo)    
                    else:
                        pass


                    #Portfolio details testing
                    try:  
                        Portfolio_tab = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_PortVal[0], sheet_name='Portfolio', header=header_no)
                        Portfolio_costunr = Portfolio_tab['Cost Unrealized Fair Value Recalculation']
                        Portfolio_fairval = Portfolio_tab['Fair Value Recalculation']
                    except:
                        PortfolioTabNotFound = True

                    if (abs(Portfolio_costunr)>1).any() or (abs(Portfolio_fairval)>1).any():
                        os.chdir(output_path+'/'+screenshot_path)
                        workbook.Sheets('Portfolio').Select()
                        workbook.Sheets('Portfolio').Cells(1,1).Select()
                        workbook.Sheets('Portfolio').UsedRange.Find('Cost Unrealized Fair Value Recalculation').Select()
                        pyautogui.hotkey('ctrl', 'down')
                        pyautogui.press('down',presses=12)
                        time.sleep(1)
                        CostUnrNotZero = pyautogui.screenshot('CostUnrNotZero.png',region=(814,500,500,262))#Edited egion=(10,305, 1700, 600))
                        CostUnrNotZero = add_border(CostUnrNotZero,'CostUnrNotZero.png')

                        if (abs(Portfolio_costunr)>1).any() and (abs(Portfolio_fairval)>1).any():
                            sec3 = 'There are differences in the "Cost Unrealized Fair Value Recalculation" and "Fair Value Recalculation" columns in the “Portfolio” tab of Portfolio working paper.'
                            reco_port = 'The audit team shall investigate the portfolio input file and determine the reason for the variance.'
                        elif (abs(Portfolio_costunr)>1).any():
                            sec3 = 'There are differences in the "Cost Unrealized Fair Value Recalculation" column in the “Portfolio” tab of Portfolio working paper.'
                            reco_port = 'The audit team shall investigate the portfolio input file and determine the reason for the variance.'
                        elif (abs(Portfolio_fairval)>1).any():
                            sec3 = 'There are differences in the "Fair Value Recalculation" column in the “Portfolio” tab of Portfolio working paper.'
                            reco_port = 'The audit team shall investigate the portfolio input file and determine the reason for the variance.'
                        else:
                            pass

                        memo_final(memo_text = sec3,screenshot = [('CostUnrNotZero.png',Inches(2.5))],reco_text = reco_port)
                    else:
                        pass

                    workbook.Close(SaveChanges=False)
                    #xl.Application.Quit()

                #####Realized Gain/Loss Testing ##############################

                if len(WAMapps_RGL)!=0:
                    xl=client.gencache.EnsureDispatch("Excel.Application")
                    xl.Visible = True
                    xl.WindowState = client.constants.xlMaximized
                    xl.FormulaBarHeight = 2
                    workbook = xl.Workbooks.Open(output_path+'/'+wp_path+'/'+WAMapps_RGL[0])
                    workbook.Save()

                    try:
                        RGL_Summary = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_RGL[0], sheet_name='Summary', header=header_no+1, skipfooter=1)
                        RGL_Detail = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_RGL[0], sheet_name='Realized_Gain_Loss', header=header_no)
                        RGL_temp0 = RGL_Summary.tail(2)
                        RGL_initialDiff = RGL_temp0['Unnamed: 2'].head(1).iloc[0]
                        RGL_UB1 = RGL_temp0['Unnamed: 2'].tail(1).fillna(0).iloc[0]
                        RGL_diff = RGL_initialDiff+RGL_UB1
                        ##for step 2 purposes
                        RGL_input = RGL_Summary[RGL_Summary['Unnamed: 1']=='Value per Realized_Gain_Loss tab']['Unnamed: 2'].fillna(0).iloc[0]
                    except:
                        RGL_NotFound = True

                    if abs(RGL_diff)>1:
                        os.chdir(output_path+'/'+screenshot_path)
                        #navigate thru the workbook
                        workbook.Sheets('Summary').Select()
                        xl.ActiveWindow.Zoom = 70
                        workbook.Sheets('Summary').Cells(1,1).Select()
                        time.sleep(1)
                        RGLNotZero = pyautogui.screenshot('RGL.png',region=(215, 512, 554, 168))# Edited region=(200,500, 600, 200))
                        RGLNotZero = add_border(RGLNotZero,'RGL.png')

                        #recommendationd
                        rgl_reco =[]
                        reco_rgl = []
                        rgl_reclass_tab = []
                        rgl_inputs = [(RGL_diff,dummy,'RGL','RGL','RGL_Reclass',Leads_Sheets,'Coversheet',rgl_reco,rgl_reclass_tab,reco_rgl)]
                        recommendation(rgl_inputs)

                        #memo
                        memo_final(memo_text ='Value per RGL tab in the RGL working paper does not reconcile with Z1 Leads.',screenshot = [('RGL.png',Inches(3))],reco_text=''.join(rgl_reco))    
                        rgl_inputs = [(reco_rgl,rgl_reclass_tab,'RGL')]
                        refer_memo(rgl_inputs)        


                    #RGL Recalculation
                    try:
                        RGL_tab = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_RGL[0], sheet_name='Realized_Gain_Loss', header=header_no)
                        RGL_basereal = RGL_tab['Base Realized Recalculated Difference']
                        RGL_diffinreal = RGL_tab['Difference in Realized']
                    except:
                        RGLDetailNotFound = True

                    if (abs(RGL_basereal)>1).any() or (abs(RGL_diffinreal)>1).any():
                        os.chdir(output_path+'/'+screenshot_path)
                        workbook.Sheets('Realized_Gain_Loss').Select()
                        workbook.Sheets('Realized_Gain_Loss').UsedRange.Find('Difference in Realized').Select()
                        # pyautogui.hotkey('alt', 'down')
                        # if (RGL_diffinreal[(RGL_tab['Difference in Realized'])<-1].count()) > 1:
                        #     pyautogui.press('down', presses=2)
                        # else:
                        #     pyautogui.press('down')

                        # pyautogui.press('enter')
                        pyautogui.hotkey('ctrl', 'down')
                        pyautogui.press('down',presses=13)
                        time.sleep(1)

                        RGLNotZero = pyautogui.screenshot('RGLNotZero.png',region=(210,420,892,400))#Edited region=(10,305, 1700, 600))
                        RGLNotZero = add_border(RGLNotZero,'RGLNotZero.png')

                        if (abs(RGL_basereal)>1).any() and (abs(RGL_diffinreal)>1).any():
                            sec3 = 'There are differences in the "Base Realized Recalculated Difference" and "Difference in Realized" columns in the “Realized_Gain_Loss” tab of Realized Gain Loss working paper.'
                            reco_rgl = 'The audit team shall investigate the RGL input file and determine the reason for the variance.'
                        elif (abs(RGL_basereal)>1).any():
                            sec3 = 'There are differences in the "Base Realized Recalculated Difference" column in the “Realized_Gain_Loss” tab of Realized Gain Loss working paper.'
                            reco_rgl = 'The audit team shall investigate the RGL input file and determine the reason for the variance.'
                        elif (abs(RGL_diffinreal)>1).any():
                            sec3 = 'There are differences in the "Difference in Realized" column in the “Realized_Gain_Loss” tab of Realized Gain Loss working paper.'
                            reco_rgl = 'The audit team shall investigate the RGL input file and determine the reason for the variance.'
                        else:
                            pass

                        memo_final(memo_text = sec3,screenshot = [('RGLNotZero.png',Inches(5))],reco_text=reco_rgl)
                    else:
                        pass


                    ###missing instrument/transaction mapping testing
                    
                    if (RGL_Detail['Instrument Sub Class'].isnull()).all():
                        pass
                    else:
                        if(RGL_Detail[(RGL_Detail['Instrument Sub Class'].str.lower()=='missing')]['Fund ID'].notnull()).any() or (RGL_Detail[(RGL_Detail['Transaction Type'].str.lower()=='missing')]['Fund ID'].notnull()).any():
                            os.chdir(output_path+'/'+screenshot_path)
                            #navigate thru the workbook
                    
                            workbook.Sheets('Realized_Gain_Loss').Select()
                            workbook.Sheets('Realized_Gain_Loss').Cells(1,1).Select()
                            xl.ActiveWindow.Zoom = 60
                            try:
                                with trial: workbook.Sheets('Realized_Gain_Loss').Range('K:K').Find('Missing').Select()
                            except:
                                with trial: workbook.Sheets('Realized_Gain_Loss').Range('K:K').Find('MISSING').Select()
                            try:
                                with trial: workbook.Sheets('Realized_Gain_Loss').Range('L:L').Find('Missing').Select()
                            except:
                                with trial: workbook.Sheets('Realized_Gain_Loss').Range('L:L').Find('MISSING').Select()
                            time.sleep(1)
                            RGLNotZero = pyautogui.screenshot('RGLMissingInsTran.png',region=(72, 420, 1734, 549))
                            RGLNotZero = add_border(RGLNotZero,'RGLMissingInsTran.png')

                            memo_final(memo_text ='There are transactions with missing instrument/transaction mapping in the RGL workpaper.',screenshot = [('RGLMissingInsTran.png',Inches(6))],reco_text='WAMapps Delivery Team should inquire this to audit team and rerun.')
                        else:
                            pass


                    workbook.Close(SaveChanges=False)
                    #xl.Application.Quit()

                ####UNSETTLED TRADES TESTING
                if len(WAMapps_Trades)!=0:
                    xl=client.gencache.EnsureDispatch("Excel.Application")
                    xl.Visible = True
                    xl.WindowState = client.constants.xlMaximized
                    xl.FormulaBarHeight = 2
                    workbook = xl.Workbooks.Open(output_path+'/'+wp_path+'/'+WAMapps_Trades[0])
                    workbook.Save()

                    try:
                        UT_Summary = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_Trades[0], sheet_name='Summary', header=header_no+8, skipfooter=1)
                        with trial: UT_Rec_Tab = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_Trades[0], sheet_name='Trades_Receivable', header=header_no)
                        with trial: UT_Pay_Tab = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_Trades[0], sheet_name='Trades_Payable', header=header_no)
                        with trial: Cap_Rec_Tab = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_Trades[0], sheet_name='Capital_Receivable', header=header_no)
                        with trial: Cap_Pay_Tab = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_Trades[0], sheet_name='Capital_Payable', header=header_no)
                        Unsettled_temp0 = UT_Summary[UT_Summary['Unnamed: 1']=='Difference'].fillna(0)
                        Unsettled_temp1 = UT_Summary.tail(2)
                        Unsettled_CapTab = Unsettled_temp1.head(1).fillna(0)['Unnamed: 2'].iloc[0]
                        Unsettled_CapLead = Unsettled_temp1.tail(1).fillna(0)['Unnamed: 2'].iloc[0]
                        Unsettled_Difference = Unsettled_temp0['Unnamed: 2'].round(2)
                        Unsettled_CapDiff = (Unsettled_CapTab-Unsettled_CapLead).round(2)
                        #for step2
                        TradesReceivable = UT_Summary[UT_Summary['Unnamed: 1']=='Value per Trades_Receivable tab' ]['Unnamed: 2'].fillna(0).iloc[0]
                        TradesPayable = UT_Summary[UT_Summary['Unnamed: 1']=='Value per Trades_Payable tab' ]['Unnamed: 2'].fillna(0).iloc[0]
                        CapitalReceivable = UT_Summary[UT_Summary['Unnamed: 1']=='Value per Capital_Receivable tab' ]['Unnamed: 2'].fillna(0).iloc[0]
                        CapitalPayable = UT_Summary[UT_Summary['Unnamed: 1']=='Value per Capital_Payable tab' ]['Unnamed: 2'].fillna(0).iloc[0]
                    except:
                        UT_NotFound = True

                    if (abs(Unsettled_Difference)>1).any() or abs(Unsettled_CapDiff)>1:
                        os.chdir(output_path+'/'+screenshot_path)
                        #navigate thru the workbook
                        workbook.Sheets('Summary').Select()
                        workbook.Sheets('Summary').Cells(1,1).Select()
                        xl.ActiveWindow.FreezePanes = False
                        xl.ActiveWindow.Zoom = 60
                        pyautogui.press('down', presses=50)
                        time.sleep(1)
                        UnsettledNotZero = pyautogui.screenshot('UnsettledNotZero.png',region=(107,413,415,531))#Edited region=(60,350, 500, 550))
                        UnsettledNotZero = add_border(UnsettledNotZero,'UnsettledNotZero.png')
                        # to define the differences per mapping

                        unsettled_names = []
                        trades_reco = []
                        reco_tradesrec = []
                        reco_tradespay = []
                        reco_caprec = []
                        reco_cappay = []
                        tradesrec_reclass_tab = []
                        tradespay_reclass_tab = [] 
                        caprec_reclass_tab = [] 
                        cappay_reclass_tab = []
                        unsettled_inputs = []
                        UT_Sheets = pd.ExcelFile(output_path+'/'+wp_path+'/'+WAMapps_Trades[0])
                        with trial: unsettled_inputs.append((Unsettled_Difference.iloc[0],unsettled_names,'Trades receivable','Trades_Rec','Trades_Rec_Reclass',UT_Sheets,'Trades_Receivable',trades_reco,tradesrec_reclass_tab,reco_tradesrec))
                        with trial: unsettled_inputs.append((Unsettled_Difference.iloc[1],unsettled_names,'Trades payable','Trades_Pay','Trades_Pay_Reclass',UT_Sheets,'Trades_Payable',trades_reco,tradespay_reclass_tab,reco_tradespay))
                        with trial: unsettled_inputs.append((Unsettled_Difference.iloc[2],unsettled_names,'Capital receivable','Cap_Rec','Cap_Rec_Reclass',UT_Sheets,'Capital_Receivable',trades_reco,caprec_reclass_tab,reco_caprec))
                        with trial: unsettled_inputs.append((Unsettled_CapDiff,unsettled_names,'Capital payable','Cap_Pay','Cap_Pay_Reclass',UT_Sheets,'Capital_Payable',trades_reco,cappay_reclass_tab,reco_cappay))

                        recommendation(unsettled_inputs)

                        unsettled_names_str = ', '.join(unsettled_names)
                        trades_rec_str = ' \n\n'.join(trades_reco)

                        if len(unsettled_names)>1:
                            preposition = 'do not reconcile with their Leads counterparts'
                        else:
                            preposition = 'does not reconcile with its Leads counterpart'

                        memo_final(memo_text ='%s per Unsettled trades working paper %s.'%(unsettled_names_str,preposition),screenshot = [('UnsettledNotZero.png',Inches(2.5))],reco_text=trades_rec_str)    
                        unsettled_inputs = [(reco_tradesrec,tradesrec_reclass_tab,'trades receivable'),(reco_tradespay,tradespay_reclass_tab,'trades payable'),(reco_caprec,caprec_reclass_tab,'capital receivable'),(reco_cappay,cappay_reclass_tab,'capital payable')]
                        refer_memo(unsettled_inputs)

                    workbook.Close(SaveChanges=False)    
                    #xl.Application.Quit()

                ################NAV BASED FEES WORKPAPER TESTING#######################################################

                if len(WAMapps_NAVbased)!=0:
                    try:
                        xl=client.gencache.EnsureDispatch("Excel.Application")
                        xl.Visible = True
                        xl.WindowState = client.constants.xlMaximized
                        xl.FormulaBarHeight = 2
                        #xl.ActiveWindow.Zoom = 70
                        xl.DisplayAlerts = True
                        workbook = xl.Workbooks.Open(output_path+'/'+wp_path+'/'+WAMapps_NAVbased[0])
                        workbook.CheckCompatibility = True
                        workbook.Save()

                    except:
                        pass

                    try:
                        NAV_shares = pd.read_excel(output_path+'/'+wp_path+'/'+WAMapps_NAVbased[0], sheet_name ='Fund_Level_Basic')
                        startIndex = NAV_shares[NAV_shares['Unnamed: 3']=='NAV'].index[0]
                        endIndex = NAV_shares[NAV_shares['Unnamed: 3']=='Total Fee:'].index[0]
                        NAV_values = NAV_shares.iloc[startIndex+1:endIndex]
                        NAV_issue = NAV_values[NAV_values['Unnamed: 3']<=0]
                        NAV_periods = NAV_issue['Unnamed: 1'].tolist()
                        NAV_reco = 'Periods %s have NAV values that are less than or equal to zero.'%(','.join(NAV_periods))
                    except:
                        NAVnotFound = True

                    try:
                        if (NAV_values['Unnamed: 3']<=0).any():
                            sec3 = docMemo.tables[1].cell(3,2).add_paragraph(style='List Bullet')
                            sec3.add_run().add_text(NAV_reco)
                            docMemo.tables[1].cell(3,2).add_paragraph().add_run().add_break()
                            p16 = docMemo.tables[2].cell(1,0).add_paragraph(style='List Bullet')
                            r16 = p16.add_run()
                            r16.add_text(NAV_reco)
                            x16 = docMemo.tables[2].cell(1,0).add_paragraph()
                            x16.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
                            #x16.add_run().add_picture('UnsettledNotZero.png',width=Inches(2.5))
                            docMemo.tables[2].cell(1,0).add_paragraph(style='List Paragraph').add_run('Recommendation:')
                            reco16 = docMemo.tables[2].cell(1,0).add_paragraph(style='List Paragraph').add_run('The audit team shall investigate the NAV input files and confirm the root cause of the variance.')
                            reco16.bold = False
                            reco16.add_break()
                    except:
                        pass


                    try:
                        workbook.Close(SaveChanges=False) 
                    except:
                        pass
                    #xl.Application.Quit()

                #If no issue on the workpapers        
                wp_issues = []
                for para in docMemo.tables[2].cell(1,0).paragraphs:
                    wp_issues.append(para.text)

                for i in range(1, len(wp_issues)):
                    if wp_issues[i]!='':
                        wp_issues_final ='With Issues'

            #     if wp_issues_final == 'With Issues':
            #         docMemo.tables[1].cell(3,2).add_paragraph().add_run().add_text('See below for more details of the issues noted.')
            #     else:
            #         docMemo.tables[1].cell(3,2).add_paragraph().add_run().add_text('No issues noted.')
                if issues_counter != 0:
                    docMemo.tables[1].cell(3,2).add_paragraph().add_run().add_text('See below for more details of the issues noted.')
                else:
                    docMemo.tables[1].cell(3,2).add_paragraph().add_run().add_text('No issues noted.')

                ###STEP 2 - Determining the completeness of input files
                #######################SCRIPT TO CHECK THE REQUIRED INPUTS OF SPECIFIC TPA ##################################################
                table = docMemo.tables[-1]

                # Data will be a list of rows represented as dictionaries
                # containing each row's data.
                data = []

                keys = None
                for i, row in enumerate(table.rows):
                    text = (cell.text for cell in row.cells)

                    # Establish the mapping based on the first row
                    # headers; these will become the keys of our dictionary
                    if i == 0:
                        keys = tuple(text)
                        continue

                    # Construct a dictionary for this row, mapping
                    # keys to values for this row
                    row_data = dict(zip(keys, text))
                    data.append(row_data)

                ##to generate output list of input files
                output_list=[]
                for i in range(0,len(data)):
                    x = data[i][tuple(data[i])[1]]##['Functionalities']
                    output_list.append(x)

                #convert to lowercase list items
                output_list = [x.lower() for x in output_list]

                ###to define the required outputs per TPA based on word template
                req_leads = [file for file in output_list if "lead sheets" in file]
                req_capact = [file for file in output_list if "capital activity" in file]
                req_cash = [file for file in output_list if "cash reconciliation" in file]
                req_dividends = [file for file in output_list if "dividend" in file]
                req_interest = [file for file in output_list if "interest" in file]
                req_IRF = [file for file in output_list if "investment roll" in file]
                req_portVal = [file for file in output_list if "portfolio val" in file]
                req_RGL = [file for file in output_list if "realized gain loss" in file]
                req_trades = [file for file in output_list if "unsettled trades" in file]
                req_FX =[file for file in output_list if "fx reasonability" in file]
                req_NAVcalc = [file for file in output_list if  "nav based fees" in file]

                missing_files = []
                missing_wp = []
                missing_tab = []

                if len(req_leads)!=0:
                    if len(WAMapps_Leads)!=0:
                        if (Leads_TB['Current_Year']==0).all():
                            missing_files.append('CY Trial Balance')
                        if (Leads_TB['Previous_Year']==0).all():
                            missing_files.append('PY Trial Balance')
                            missing_tab.append('PY Trial Balance')
                        #beta missing
                        if 'Account_Rec' not in Leads_Sheets.sheet_names and TPA =='beta':
                            missing_files.append('CY FRAN')
                            missing_wp.append('Adjustments, Account_Rec, Balance_Sheet_FS, and Profit_Loss_FS tabs of the Leads')
                    else:
                        missing_files.append('PY & CY Trial Balance')
                        missing_wp.append('Leads')

                if len(req_capact)!=0:
                    if len(WAMapps_CapAct)!=0:
                        if ((SC_T1['Description']=='Opening Net Assets per System')&(SC_T1['Base Amount'].isna())).any():
                            missing_files.append('PY Share Register')
                            missing_tab.append('PY Share Register')
                        elif ((SC_T1['Description']=='Closing Net Assets per System')&(SC_T1['Base Amount'].isna())).any():
                            missing_files.append('CY Share Register')
                        else:
                            pass

                        if (CA_Detail['Fund ID'] !=str(doc_fr['FundShortCode'][row_num])).all():
                            missing_files.append('CY Capital Activity')
                            missing_wp.append('Capital Activity')
                        else:
                            pass
                    else:
                        missing_files.append('PY & CY Share Register')
                        missing_files.append('CY Capital Activity')
                        missing_wp.append('Capital Activity')


                if len(req_cash)!=0:
                    if len(WAMapps_CashRecon)!=0:
                        if (Cash_detail['Fund ID']!=doc_fr['FundShortCode'][row_num]).all():
                            missing_files.append('CY Cash Listing')
                            missing_wp.append('Cash')
                    else:
                        missing_files.append('CY Cash Listing')
                        missing_wp.append('Cash')

                if len(req_dividends)!=0:
                    if len(WAMapps_Dividends)!=0:
                        Dividends_Sheets = pd.ExcelFile(output_path+'/'+wp_path+'/'+WAMapps_Dividends[0])
                        Div_Inc_Tab_Missing = False if 'Dividends_Income' in Dividends_Sheets.sheet_names else True
                        Div_Expense_Tab_Missing = False if 'Dividends_Expense' in Dividends_Sheets.sheet_names else True
                        Div_Rec_Tab_Missing = False if 'Dividends_Receivable' in Dividends_Sheets.sheet_names else True
                        Div_Pay_Tab_Missing = False if 'Dividends_Payable' in Dividends_Sheets.sheet_names else True

                        if Div_Inc_Tab_Missing and Div_Expense_Tab_Missing and Div_Rec_Tab_Missing and Div_Pay_Tab_Missing:
                            missing_files.append('CY Dividends')
                            missing_wp.append('Dividends')
                        elif Div_Inc_Tab_Missing and Div_Expense_Tab_Missing:
                            missing_files.append('CY Dividends Income/Expense')
                            missing_tab.append('CY Dividends income/Expense')
                        else:
                            pass


                        if Div_Inc_Tab_Missing and Div_Expense_Tab_Missing and Div_Rec_Tab_Missing and Div_Pay_Tab_Missing:
                            missing_files.append('CY Dividends')
                            missing_wp.append('Dividends')
                        elif Div_Rec_Tab_Missing and Div_Pay_Tab_Missing:
                            missing_files.append('CY Dividends Receivable/Payable')
                            missing_tab.append('CY Dividends Receivable/Payable')
                        else:
                            pass
                    else:
                        missing_files.append('CY Dividends')
                        missing_wp.append('Dividends')

                if len(req_interest)!=0:
                    if len(WAMapps_Interest)!=0:
                        Interest_Sheets = pd.ExcelFile(output_path+'/'+wp_path+'/'+WAMapps_Interest[0])
                        Int_Inc_Tab_Missing = False if 'Interest_Income' in Interest_Sheets.sheet_names else True
                        Int_Expense_Tab_Missing = False if 'Interest_Expense' in Interest_Sheets.sheet_names else True
                        Int_Rec_Tab_Missing = False if 'Interest_Receivable' in Interest_Sheets.sheet_names else True
                        Int_Pay_Tab_Missing = False if 'Interest_Payable' in Interest_Sheets.sheet_names else True

                        if Int_Inc_Tab_Missing and Int_Expense_Tab_Missing and Int_Rec_Tab_Missing and Int_Pay_Tab_Missing:
                            missing_files.append('CY Interest')
                            missing_wp.append('Interest')
                        elif Int_Inc_Tab_Missing and Int_Expense_Tab_Missing:
                            missing_files.append('CY Interest Income/Expense')
                            missing_tab.append('CY Interest Income/Expense')
                        else:
                            pass


                        if Int_Inc_Tab_Missing and Int_Expense_Tab_Missing and Int_Rec_Tab_Missing and Int_Pay_Tab_Missing:
                            missing_files.append('CY Interest')
                            missing_wp.append('Interest')
                        elif Int_Rec_Tab_Missing and Int_Pay_Tab_Missing:
                            missing_files.append('CY Interest Receivable/Payable')
                            missing_tab.append('CY Interest Receivable/Payable')
                        else:
                            pass
                    else:
                        missing_files.append('CY Interest')
                        missing_wp.append('Interest')

                if len(req_IRF)!=0:
                    if len(WAMapps_InvRF)!=0:
                        if (OpeningFV==0 or type(OpeningFV)!=np.float64) and (abs(IRF_Difference.iloc[1])>1):
                            missing_files.append('PY Portfolio')
                            missing_tab.append('PY Portfolio')
                        if (OpeningCost==0 or type(OpeningCost)!=np.float64) and (abs(IRF_PyDiff)>1):
                            missing_files.append('PY Portfolio')
                            missing_tab.append('PY Portfolio')
                    else:
                        missing_files.append('PY Portfolio')
                        missing_files.append('CY Purchases & Sales')

                if len(req_portVal)!=0:
                    if len(WAMapps_PortVal)!=0:
                        if (PortVal_TietoD1['Original Fair Value per Portfolio'].isna()).all():
                            missing_files.append('CY Portfolio')
                    else:
                        missing_files.append('CY Portfolio')
                        missing_wp.append('Portfolio')

                if len(req_RGL)!=0:
                    if len(WAMapps_RGL)!=0:
                        if (RGL_Detail['Fund ID']!=doc_fr['FundShortCode'][row_num]).all():
                            missing_files.append('CY RGL')
                            missing_wp.append('RGL')
                    else:
                        missing_files.append('CY RGL')
                        missing_wp.append('RGL')


                if len(req_trades)!=0:
                    if len(WAMapps_Trades)!=0:
                        UT_Sheets = pd.ExcelFile(output_path+'/'+wp_path+'/'+WAMapps_Trades[0])
                        UT_Rec_Tab_Missing = False if 'Trades_Receivable' in UT_Sheets.sheet_names else True
                        UT_Pay_Tab_Missing = False if 'Trades_Payable' in UT_Sheets.sheet_names else True
                        Cap_Rec_Tab_Missing = False if 'Capital_Receivable' in UT_Sheets.sheet_names else True
                        Cap_Pay_Tab_Missing = False if 'Capital_Payable' in UT_Sheets.sheet_names else True

                        if UT_Rec_Tab_Missing and UT_Pay_Tab_Missing and Cap_Rec_Tab_Missing and Cap_Pay_Tab_Missing:
                            missing_files.append('CY Open Trades')
                            missing_wp.append('Open Trades')
                        elif UT_Rec_Tab_Missing and UT_Pay_Tab_Missing:
                            missing_files.append('CY Open Trades Receivable/Payable')
                            missing_tab.append('CY Open Trades Receivable/Payable')
                        else:
                            pass


                        if UT_Rec_Tab_Missing and UT_Pay_Tab_Missing and Cap_Rec_Tab_Missing and Cap_Pay_Tab_Missing:
                            missing_files.append('CY Open Trades')
                            missing_wp.append('Open Trades')
                        elif Cap_Rec_Tab_Missing and Cap_Pay_Tab_Missing:
                            missing_files.append('CY Capital Receivable/Payable')
                            missing_tab.append('CY Capital Receivable/Payable')
                        else:
                            pass
                    else:
                        missing_files.append('CY Open Trades')
                        missing_wp.append('Open Trades')


                if len(req_NAVcalc)!=0:
                    if len(WAMapps_NAVbased)==0:
                        missing_files.append('CY NAV')


                missing_files = list(set(missing_files))
                #missing_files_str = ', '.join(missing_files)
                missing_files_str = '\n- '.join(missing_files)

                missing_wp = list(set(missing_wp))
                #For proper grammar
                if len(missing_wp) > 1:
                    missing_wp[-1] = 'and ' + missing_wp[-1]
                else:
                    pass

                if len(missing_wp) > 2:
                    missing_wp_str = ', '.join(missing_wp)
                else:
                    missing_wp_str = ' '.join(missing_wp)

                missing_tab = list(set(missing_tab))
                #For proper grammar
                if len(missing_tab) > 1:
                    missing_tab[-1] = 'and ' + missing_tab[-1]
                else:
                    pass

                if len(missing_tab) > 2:
                    missing_tab_str = ', '.join(missing_tab)
                else:
                    missing_tab_str = ' '.join(missing_tab)

                #Completeness W/ Recommendation
                if len(missing_files)!=0:
                    if len(missing_tab)!=0 and len(missing_wp)==0:
                        sec2a = 'No files received or no transactions noted for:\n- %s\n\nThe %s balances in the affected working papers will be nil.'%(missing_files_str,missing_tab_str)
                        sec3b = 'The audit team shall confirm whether the files are really missing.'   
                    elif len(missing_tab)!=0:
                        sec2a = 'No files received or no transactions noted for:\n- %s\n\nAs such, the WAMapps %s working papers will not be produced. Also, the %s balances in the affected working papers will be nil.'%(missing_files_str, missing_wp_str, missing_tab_str)
                        sec3b = 'The audit team shall confirm whether there are really no transactions regarding the above accounts or whether the files are really missing.'
                    else:
                        sec2a = 'No files received or no transactions noted for:\n- %s\n\nAs such, the WAMapps %s working papers will not be produced.'%(missing_files_str,missing_wp_str)
                        sec3b = 'The audit team shall confirm whether there are really no transactions regarding the above accounts or whether the files are really missing.'

                    completeness(memo_text = sec2a, reco_text = sec3b)

                else:
                    docMemo.tables[1].cell(2,2).add_paragraph().add_run().add_text('No issues noted.')
                    docMemo.tables[2].cell(1,0).paragraphs[1].add_run().add_break()
                    docMemo.tables[2].cell(1,0).paragraphs[1].add_run().add_text('No issues noted.')

                ##defining mapping columns for section 1

                account_mapping_cols = ["tpa",
                "join_1",
                "join_2",
                "join_3",
                "join_4",
                "description",
                "4_col_join",
                "account_type",
                "account_sub_type",
                "account_class",
                "account_lead",
                "long_short",
                "cost_unrealized",
                "fs_type",
                "fs_class",
                "cf_fihi_class",
                "capital_income",
                "custom_flag_1",
                "custom_flag_2"]

                currency_mapping_cols = [
                    "tpa","currency_description","currency_code"
                ]

                instrument_mapping_cols = ["tpa","join_1","join_2","description","instrument_class","instrument_sub_class","exclude","mv_unr",
                                        "custom_flag_1","custom_flag_2"]

                share_mapping_cols = ["tpa","seq_1","join_1","seq_2","join_2","sc_sub_type","sc_type","soc_type","multiplier","exclude"]

                transaction_mapping_cols = ["tpa","join_1","join_2","description","transaction_class","transaction_sub_class","transaction_class_2","multiplier","exclude"]

                os.chdir(output_path+'/'+'Workpapers')

                ###section 1 mapping auto attach
                mapping_files = [file for file in files if "csv" in file[-3:]]
                mapping_text = []
                custom_mapping = []

                #to loop thru the mapping files:
                for mapping_type in mapping_files:
                    map1 = pd.read_csv(mapping_type)
                    map_cols = map1.columns
                    account_match = len(list(set(map_cols).intersection(account_mapping_cols)))/len(account_mapping_cols)
                    curr_match = len(list(set(map_cols).intersection(currency_mapping_cols)))/len(currency_mapping_cols)
                    instr_match = len(list(set(map_cols).intersection(instrument_mapping_cols)))/len(instrument_mapping_cols)
                    share_match = len(list(set(map_cols).intersection(share_mapping_cols)))/len(share_mapping_cols)
                    trans_match = len(list(set(map_cols).intersection(transaction_mapping_cols)))/len(transaction_mapping_cols)
                    if account_match>.8 and (map1['tpa'][0]==doc_fr['TPAName'][row_num]):
                        mapping_text.append('account')
                        custom_mapping.append(mapping_type)
                    elif curr_match>.8 and (map1['tpa'][0]==doc_fr['TPAName'][row_num]):
                        mapping_text.append('currency')
                        custom_mapping.append(mapping_type)
                    elif instr_match>.8 and (map1['tpa'][0]==doc_fr['TPAName'][row_num]):
                        mapping_text.append('instrument')
                        custom_mapping.append(mapping_type)
                    elif share_match>.8 and (map1['tpa'][0]==doc_fr['TPAName'][row_num]):
                        mapping_mapping_text.append('share capital')
                        custom_mapping.append(mapping_type)
                    elif trans_match>.8 and (map1['tpa'][0]==doc_fr['TPAName'][row_num]):
                        mapping_text.append('transaction')
                        custom_mapping.append(mapping_type)
                    else:
                        pass
              
                mapping_text_str = ', '.join(mapping_text)
                custom_mapping_str = ', '.join(custom_mapping)

                sec1a = docMemo.tables[1].cell(1,2).add_paragraph(style='List Bullet')
                sec1a.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                sec1a.add_run().add_text("Changes were applied on some of the input files. Please refer to the attached transformation tracker below for the list of changes made.")
                docMemo.tables[1].cell(1,2).add_paragraph().add_run().add_break()
                sec1b = docMemo.tables[1].cell(1,2).add_paragraph(style='List Bullet')
                sec1b.add_run().add_text("The following custom %s mapping were used in processing this request:" %(mapping_text_str))
                sec1b.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                docMemo.tables[1].cell(1,2).add_paragraph().add_run().add_text('\t'+custom_mapping_str.replace(', ', '\n\t'))
                docMemo.tables[1].cell(1,2).paragraphs[0].style='List Paragraph'


                ##save the transformation tracker memo
                os.chdir(output_path)
                finalMemo = 'FY%s %s (%s) WAMapps Delivery Memo P%s.docx' %(dateString,doc_fr['FundName'][row_num], doc_fr['FundShortCode'][row_num],period)
                docMemo.save(finalMemo)

                word = client.gencache.EnsureDispatch("Word.Application") # Using DispatchEx for an entirely new Word instance

                word.Visible = True
                docMemo_win32 = word.Documents.Open(output_path+'/'+finalMemo)
                word.ActiveWindow.ActivePane.View.Zoom.Percentage = 100

                #Disable save with comments warning
                word.Options.WarnBeforeSavingPrintingSendingMarkup = False


                ###to highlight to red all reclassifications done
                for name in set(account_names_append):
                    word.Selection.Find.Execute(FindText = name, MatchCase=True)
                    if word.Selection.Find.Found:
                        word.Selection.Font.ColorIndex = 6
                        word.Selection.Font.Bold = True
                        wdStory=6
                        word.Selection.HomeKey(Unit=wdStory)

                for aclass in set(account_class_append):
                    word.Selection.Find.Execute(FindText = aclass, MatchCase=True)
                    if word.Selection.Find.Found:
                        word.Selection.Font.ColorIndex = 6
                        word.Selection.Font.Bold = True
                        wdStory=6
                        word.Selection.HomeKey(Unit=wdStory)

                ###to add the mapping attachments to the memo
                #docMemo_win32.InlineShapes.AddOLEObject(ClassType = 'Excel.Sheet', FileName=output_path+'/'+transformation_file, DisplayAsIcon=True, IconFileName=transformation_file, IconIndex=2)
                #docMemo_win32.InlineShapes.AddOLEObject(FileName=output_path+'/'+transformation_file, IconFileName="wordicon.exe",  DisplayAsIcon=True, IconIndex=0, IconLabel=transformation_file, Range=docMemo_win32.Tables(2).Cell(2,3).Range)

                docMemo_win32.Save()
                docMemo_win32.Close()
                os.startfile(output_path+'/'+finalMemo)

                print("\n(FR Row Number " + str(row_num) + ") Delivery memo successfully generated.")
                print("- - - %s seconds - - -" % (round(time.time() - start_time, 2)))
                time.sleep(5)

                # #Delete all current variables to avoid bugs when batch processing different requests
                # for var in dir():
                #     if var not in Global_var:
                #         del globals()[var]
                
                #Delete temp screenshots folder
                os.chdir(output_path)
                shutil.rmtree( r'screenshots')

    print('\nTool finished running.')
  except:
    print('\nTool ran into an error: \n\n',traceback.format_exc(), '\n\n Please submit this error to Driving Digital Team to be fixed.')


root.mainloop()

# Flushing after close of  tkinter root
sys.stdout = old_stdout